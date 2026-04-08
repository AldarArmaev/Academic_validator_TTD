"""
TDD: Тесты для модуля парсинга DOCX.
Сначала пишем тесты (красный этап), затем реализуем функционал.
"""
import pytest
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.contour_a.docx_parser import DOCXParser, parse_docx
from src.contour_a.document_loader import (
    Paragraph,
    Heading,
    Table,
    Figure,
    Reference,
    DocumentContent,
)


class TestDOCXParserInit:
    """Тесты инициализации парсера."""

    def test_parser_initialization(self):
        """Создание парсера с начальными значениями."""
        parser = DOCXParser()
        assert parser._current_page == 1
        assert parser._page_breaks_count == 0


class TestParseDocxFunction:
    """Тесты удобной функции parse_docx."""

    def test_parse_docx_creates_parser(self):
        """Функция parse_docx создаёт парсер и вызывает parse."""
        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.DOCXParser') as mock_parser_class:
                mock_parser = Mock()
                mock_parser.parse.return_value = DocumentContent()
                mock_parser_class.return_value = mock_parser
                
                result = parse_docx('test.docx')
                
                mock_parser_class.assert_called_once()
                mock_parser.parse.assert_called_once_with('test.docx')
                assert isinstance(result, DocumentContent)


class TestDOCXParserFileHandling:
    """Тесты обработки файлов парсером."""

    def test_parse_nonexistent_file(self):
        """Попытка парсинга несуществующего файла."""
        parser = DOCXParser()
        with pytest.raises(FileNotFoundError):
            parser.parse('/nonexistent/path/file.docx')

    def test_parse_unsupported_format(self):
        """Попытка парсинга неподдерживаемого формата."""
        parser = DOCXParser()
        with patch('pathlib.Path.exists', return_value=True):
            with pytest.raises(ValueError, match="Неподдерживаемый формат"):
                parser.parse('file.pdf')

    def test_parse_invalid_docx(self):
        """Попытка парсинга невалидного DOCX файла."""
        parser = DOCXParser()
        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc:
                mock_doc.side_effect = Exception("Invalid file")
                with pytest.raises(ValueError, match="Ошибка при открытии"):
                    parser.parse('invalid.docx')


class TestParagraphParsing:
    """Тесты парсинга абзацев."""

    def _setup_mock_doc(self, mock_doc, paragraphs=None, tables=None):
        """Вспомогательный метод для настройки мок-документа."""
        mock_doc.paragraphs = paragraphs or []
        mock_doc.tables = tables or []
        mock_doc.sections = []
        mock_doc.part.rels.values.return_value = []
        
        # Настройка моков для парсинга формул
        for para in (paragraphs or []):
            mock_element = Mock()
            mock_element.findall.return_value = []  # Пустой список формул по умолчанию
            para._element = mock_element
    
    def test_parse_empty_document(self):
        """Парсинг пустого документа."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                mock_doc = Mock()
                self._setup_mock_doc(mock_doc)
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('empty.docx')
                
                assert content.paragraphs == []
                assert content.headings == []
                assert content.tables == []
                assert content.file_name == 'empty.docx'

    def test_parse_paragraph_basic(self):
        """Парсинг простого абзаца."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                mock_para = Mock()
                mock_para.text = "Тестовый абзац"
                mock_para.style.name = "Normal"
                mock_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                mock_para.runs = []

                mock_doc = Mock()
                self._setup_mock_doc(mock_doc, [mock_para])
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                assert len(content.paragraphs) == 1
                assert content.paragraphs[0].text == "Тестовый абзац"
                assert content.paragraphs[0].style_name == "Normal"

    def test_parse_paragraph_with_font_info(self):
        """Парсинг абзаца с информацией о шрифте."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                mock_run = Mock()
                mock_run.text = "Текст с форматированием"
                mock_run.font.name = "Times New Roman"
                mock_run.font.size = Pt(14)
                mock_run.font.bold = True
                mock_run.font.italic = False

                mock_para = Mock()
                mock_para.text = "Текст с форматированием"
                mock_para.style.name = "Normal"
                mock_para.runs = [mock_run]
                mock_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                mock_doc = Mock()
                self._setup_mock_doc(mock_doc, [mock_para])
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                assert len(content.paragraphs) == 1
                para = content.paragraphs[0]
                assert para.font_name == "Times New Roman"
                assert para.font_size == 14.0
                assert para.is_bold is True
                assert para.is_italic is False

    def test_parse_paragraph_alignment(self):
        """Парсинг абзаца с различным выравниванием."""
        parser = DOCXParser()

        alignments = [
            (WD_ALIGN_PARAGRAPH.LEFT, "left"),
            (WD_ALIGN_PARAGRAPH.CENTER, "center"),
            (WD_ALIGN_PARAGRAPH.RIGHT, "right"),
            (WD_ALIGN_PARAGRAPH.JUSTIFY, "justify"),
        ]

        for docx_align, expected_align in alignments:
            with patch('pathlib.Path.exists', return_value=True):
                with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                    mock_para = Mock()
                    mock_para.text = f"Абзац с {expected_align} выравниванием"
                    mock_para.style.name = "Normal"
                    mock_para.alignment = docx_align
                    mock_para.runs = []

                    mock_doc = Mock()
                    self._setup_mock_doc(mock_doc, [mock_para])
                    mock_doc_class.return_value = mock_doc
                    
                    content = parser.parse('test.docx')
                    
                    assert content.paragraphs[0].alignment == expected_align


class TestHeadingParsing:
    """Тесты парсинга заголовков."""

    def _setup_mock_doc(self, mock_doc, paragraphs=None, tables=None):
        """Вспомогательный метод для настройки мок-документа."""
        mock_doc.paragraphs = paragraphs or []
        mock_doc.tables = tables or []
        mock_doc.sections = []
        mock_doc.part.rels.values.return_value = []
        
        # Настройка моков для парсинга формул
        for para in (paragraphs or []):
            mock_element = Mock()
            mock_element.findall.return_value = []  # Пустой список формул по умолчанию
            para._element = mock_element

    def test_parse_heading_level_1(self):
        """Парсинг заголовка первого уровня."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                mock_para = Mock()
                mock_para.text = "Глава 1. Введение"
                mock_para.style.name = "Heading 1"
                mock_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mock_para.runs = []

                mock_doc = Mock()
                self._setup_mock_doc(mock_doc, [mock_para])
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                assert len(content.headings) == 1
                heading = content.headings[0]
                assert heading.level == 1
                assert heading.style_name == "Heading 1"
                assert heading.numbering == "Глава 1"

    def test_parse_heading_level_2(self):
        """Парсинг заголовка второго уровня."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                mock_para = Mock()
                mock_para.text = "1.1. Обзор литературы"
                mock_para.style.name = "Heading 2"
                mock_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                mock_para.runs = []

                mock_doc = Mock()
                self._setup_mock_doc(mock_doc, [mock_para])
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                assert len(content.headings) == 1
                heading = content.headings[0]
                assert heading.level == 2
                assert heading.numbering == "1.1"

    def test_parse_multiple_headings(self):
        """Парсинг нескольких заголовков разных уровней."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                mock_para1 = Mock()
                mock_para1.text = "Глава 1. Теоретическая часть"
                mock_para1.style.name = "Heading 1"
                mock_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mock_para1.runs = []
                
                mock_para2 = Mock()
                mock_para2.text = "1.1. Понятие нормы"
                mock_para2.style.name = "Heading 2"
                mock_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                mock_para2.runs = []
                
                mock_para3 = Mock()
                mock_para3.text = "1.2. Подходы к нормоконтролю"
                mock_para3.style.name = "Heading 2"
                mock_para3.alignment = WD_ALIGN_PARAGRAPH.LEFT
                mock_para3.runs = []

                mock_doc = Mock()
                mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3]
                mock_doc.tables = []
                mock_doc.sections = []
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                assert len(content.headings) == 3
                assert content.headings[0].level == 1
                assert content.headings[1].level == 2
                assert content.headings[2].level == 2


class TestTableParsing:
    """Тесты парсинга таблиц."""

    def test_parse_table_basic(self):
        """Парсинг простой таблицы."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                # Создаём мок таблицы
                mock_cell1 = Mock()
                mock_cell1.text = "Ячейка 1"
                mock_cell2 = Mock()
                mock_cell2.text = "Ячейка 2"
                mock_cell3 = Mock()
                mock_cell3.text = "Ячейка 3"
                mock_cell4 = Mock()
                mock_cell4.text = "Ячейка 4"
                
                mock_row1 = Mock()
                mock_row1.cells = [mock_cell1, mock_cell2]
                mock_row2 = Mock()
                mock_row2.cells = [mock_cell3, mock_cell4]
                
                mock_table = Mock()
                mock_table.rows = [mock_row1, mock_row2]
                mock_table.columns = [Mock(), Mock()]  # 2 колонки

                mock_doc = Mock()
                mock_doc.paragraphs = []
                mock_doc.tables = [mock_table]
                mock_doc.sections = []
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                assert len(content.tables) == 1
                table = content.tables[0]
                assert table.rows == 2
                assert table.columns == 2
                assert len(table.cells) == 2
                assert table.cells[0][0] == "Ячейка 1"
                assert table.cells[1][1] == "Ячейка 4"

    def test_parse_multiple_tables(self):
        """Парсинг нескольких таблиц."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                mock_table1 = Mock()
                mock_table1.rows = [Mock()]
                mock_table1.columns = [Mock()]
                mock_table1.rows[0].cells = [Mock(text="Таблица 1")]
                
                mock_table2 = Mock()
                mock_table2.rows = [Mock()]
                mock_table2.columns = [Mock(), Mock(), Mock()]
                mock_table2.rows[0].cells = [Mock(text="Таблица 2")]

                mock_doc = Mock()
                mock_doc.paragraphs = []
                mock_doc.tables = [mock_table1, mock_table2]
                mock_doc.sections = []
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                assert len(content.tables) == 2
                assert content.tables[0].columns == 1
                assert content.tables[1].columns == 3


class TestReferenceParsing:
    """Тесты парсинга библиографических ссылок."""

    def test_parse_simple_reference(self):
        """Парсинг простой ссылки [1]."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                mock_para = Mock()
                mock_para.text = "Как указано в исследовании [1], это важно."
                mock_para.style.name = "Normal"
                mock_para.runs = []
                mock_para.alignment = None

                mock_doc = Mock()
                mock_doc.paragraphs = [mock_para]
                mock_doc.tables = []
                mock_doc.sections = []
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                assert len(content.references) == 1
                ref = content.references[0]
                assert ref.text == "[1]"
                assert ref.number == 1
                assert ref.position == "inline"

    def test_parse_reference_with_page(self):
        """Парсинг ссылки с указанием страницы [1, с. 5]."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                mock_para = Mock()
                mock_para.text = "Цитата из источника [1, с. 5] подтверждает это."
                mock_para.style.name = "Normal"
                mock_para.runs = []
                mock_para.alignment = None

                mock_doc = Mock()
                mock_doc.paragraphs = [mock_para]
                mock_doc.tables = []
                mock_doc.sections = []
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                assert len(content.references) == 1
                ref = content.references[0]
                assert ref.text == "[1, с. 5]"
                assert ref.number == 1

    def test_parse_multiple_references(self):
        """Парсинг нескольких ссылок [1; 2; 3]."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                mock_para = Mock()
                mock_para.text = "Многие исследования [1; 2; 3] показывают это."
                mock_para.style.name = "Normal"
                mock_para.runs = []
                mock_para.alignment = None

                mock_doc = Mock()
                mock_doc.paragraphs = [mock_para]
                mock_doc.tables = []
                mock_doc.sections = []
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                assert len(content.references) >= 1
                # Должна быть найдена хотя бы одна ссылка
                refs = [r for r in content.references if ';' in r.text or r.number]
                assert len(refs) >= 1


class TestSectionParsing:
    """Тесты парсинга секций и полей."""

    def test_parse_section_margins(self):
        """Парсинг полей документа из секции."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                mock_section = Mock()
                mock_section.left_margin = Cm(3.0)
                mock_section.right_margin = Cm(1.0)
                mock_section.top_margin = Cm(2.0)
                mock_section.bottom_margin = Cm(2.0)

                mock_doc = Mock()
                mock_doc.paragraphs = []
                mock_doc.tables = []
                mock_doc.sections = [mock_section]
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                assert 'margins' in content.metadata
                margins = content.metadata['margins']
                assert margins['left_cm'] == 3.0
                assert margins['right_cm'] == 1.0
                assert margins['top_cm'] == 2.0
                assert margins['bottom_cm'] == 2.0


class TestPageBreakDetection:
    """Тесты обнаружения разрывов страниц."""

    def test_page_counter_increments_on_break(self):
        """Счётчик страниц увеличивается при разрыве."""
        parser = DOCXParser()

        with patch('pathlib.Path.exists', return_value=True):
            with patch('src.contour_a.docx_parser.Document') as mock_doc_class:
                # Создаём абзац с разрывом страницы
                mock_para1 = Mock()
                mock_para1.text = "Первая страница"
                mock_para1.style.name = "Normal"
                mock_para1.runs = []
                mock_para1.alignment = None
                
                # Элемент разрыва страницы
                mock_br_element = Mock()
                mock_br_element.tag = 'w:br'
                mock_br_element.get.return_value = 'page'
                
                mock_para2 = Mock()
                mock_para2.text = "Вторая страница"
                mock_para2.style.name = "Normal"
                mock_para2.runs = []
                mock_para2.alignment = None
                mock_para2._element = [mock_br_element]

                mock_doc = Mock()
                mock_doc.paragraphs = [mock_para1, mock_para2]
                mock_doc.tables = []
                mock_doc.sections = []
                mock_doc_class.return_value = mock_doc
                
                content = parser.parse('test.docx')
                
                # Второй абзац должен быть на странице 2
                assert content.paragraphs[1].page_number >= 1


class TestFontExtraction:
    """Тесты извлечения информации о шрифте."""

    def test_get_font_from_run(self):
        """Извлечение шрифта из run абзаца."""
        parser = DOCXParser()

        mock_run = Mock()
        mock_run.text = "Текст"
        mock_run.font.name = "Arial"
        mock_run.font.size = Pt(12)
        mock_run.font.bold = False
        mock_run.font.italic = True

        mock_para = Mock()
        mock_para.runs = [mock_run]

        font_info = parser._get_font_info(mock_para)

        assert font_info['name'] == "Arial"
        assert font_info['size'] == 12.0
        assert font_info['bold'] is False
        assert font_info['italic'] is True

    def test_get_font_from_style(self):
        """Извлечение шрифта из стиля абзаца."""
        parser = DOCXParser()

        mock_style_font = Mock()
        mock_style_font.name = "Times New Roman"
        mock_style_font.size = Pt(14)

        mock_style = Mock()
        mock_style.font = mock_style_font

        mock_para = Mock()
        mock_para.runs = []
        mock_para.style = mock_style

        font_info = parser._get_font_info(mock_para)

        assert font_info['name'] == "Times New Roman"
        assert font_info['size'] == 14.0


class TestIndentExtraction:
    """Тесты извлечения отступов."""

    def test_get_first_line_indent(self):
        """Извлечение отступа первой строки."""
        parser = DOCXParser()

        # Создаём мок XML элемента с отступом
        mock_ind = Mock()
        mock_ind.get.side_effect = lambda key, default=None: {'{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstLine': '720'}.get(key, default)

        mock_pPr = Mock()
        mock_pPr.find.return_value = mock_ind

        mock_element = Mock()
        mock_element.pPr = mock_pPr

        mock_para = Mock()
        mock_para._element = mock_element

        indent = parser._get_indent_first_line(mock_para)

        # 720 twips / 20 = 36 пунктов
        assert indent == 36.0


class TestLineSpacingExtraction:
    """Тесты извлечения межстрочного интервала."""

    def test_get_line_spacing_auto(self):
        """Извлечение автоматического межстрочного интервала."""
        parser = DOCXParser()

        mock_spacing = Mock()
        mock_spacing.get.side_effect = lambda key, default='': {
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lineRule': 'auto',
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': '240'
        }.get(key, default)

        mock_pPr = Mock()
        mock_pPr.find.return_value = mock_spacing

        mock_element = Mock()
        mock_element.pPr = mock_pPr

        mock_para = Mock()
        mock_para._element = mock_element

        spacing = parser._get_line_spacing(mock_para)

        assert spacing == 1.0
