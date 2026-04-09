# tests/conftest.py
import pytest
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import json


# ── Вспомогательные функции ──────────────────────────────────────────────────

def set_paragraph_spacing(para, line_twips: int, space_before: int = 0, space_after: int = 0):
    """Устанавливает межстрочный интервал и отступы до/после абзаца через XML."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(line_twips))
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:before'), str(space_before))
    spacing.set(qn('w:after'), str(space_after))
    pPr.append(spacing)


def set_first_line_indent(para, dxa: int):
    """Устанавливает отступ первой строки через XML."""
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLine'), str(dxa))
    pPr.append(ind)


def set_alignment(para, alignment: str):
    """alignment: 'both' | 'left' | 'center' | 'right'"""
    pPr = para._p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), alignment)
    pPr.append(jc)


def add_correct_paragraph(doc, text: str):
    """Добавляет абзац со всеми правильными параметрами форматирования."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    set_paragraph_spacing(p, line_twips=420, space_before=0, space_after=0)
    set_first_line_indent(p, dxa=720)
    set_alignment(p, 'both')
    return p


def make_base_doc():
    """Создаёт документ с правильными полями."""
    doc = Document()
    s = doc.sections[0]
    EMU_PER_DXA = 635
    s.left_margin   = 1701 * EMU_PER_DXA   # 3.0 см = 1701 DXA = 1079135 EMU
    s.right_margin  = 567  * EMU_PER_DXA   # 1.0 см
    s.top_margin    = 1134 * EMU_PER_DXA   # 2.0 см
    s.bottom_margin = 1134 * EMU_PER_DXA   # 2.0 см
    return doc


# ── Фикстуры для тестов валидации шрифта ─────────────────────────────────────

@pytest.fixture(scope="session")
def correct_docx(tmp_path_factory):
    """Полностью корректный документ — 0 ошибок форматирования."""
    path = tmp_path_factory.mktemp("fix") / "correct.docx"
    doc = make_base_doc()
    for title in ["Введение", "Глава 1. Теоретические основы", "Заключение", "Список литературы"]:
        doc.add_heading(title, level=1)
        add_correct_paragraph(doc, "Текст раздела. Содержательный абзац.")
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_font_docx(tmp_path_factory):
    """Абзац с Arial вместо Times New Roman (нарушение Ф-1)."""
    path = tmp_path_factory.mktemp("fix") / "wrong_font.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    add_correct_paragraph(doc, "Правильный абзац.")
    p = doc.add_paragraph()
    run = p.add_run("Абзац с неправильным шрифтом Arial.")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    set_paragraph_spacing(p, 420)
    set_first_line_indent(p, 720)
    set_alignment(p, 'both')
    doc.add_heading("Заключение", level=1)
    doc.add_heading("Список литературы", level=1)
    doc.save(path)
    return path


@pytest.fixture
def rules():
    return {
        "font": {"family": "Times New Roman", "size_half_points": 28, "size_pt": 14},
        "paragraph": {
            "line_spacing_twips": 420,
            "line_spacing_rule": "auto",
            "first_line_indent_dxa": 720,
            "first_line_indent_cm": 1.25,
            "alignment": "both",
            "space_before_twips": 0,
            "space_after_twips": 0
        },
        "margins_dxa": {"left": 1701, "right": 567, "top": 1134, "bottom": 1134},
        "required_sections": ["введение", "заключение", "список литературы"],
        "chapter_heading_pattern": "^Глава \\d+\\.\\s.+",
        "tolerances": {"dxa": 20, "pt": 0.5},
        "references": {"min_sources": 40}
    }


# ── Фикстуры для тестов ссылок и списка литературы (Л-*) ─────────────────────

def add_reference_entry(doc, text: str):
    """Добавляет запись в список литературы."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    set_paragraph_spacing(p, line_twips=420, space_before=0, space_after=0)
    set_alignment(p, 'left')
    return p


@pytest.fixture(scope="session")
def wrong_L_4_alphabetical_order_docx(tmp_path_factory):
    """Нарушение алфавитного порядка в списке литературы (Л-4)."""
    path = tmp_path_factory.mktemp("fix") / "wrong_L_4_alphabetical.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    doc.add_heading("Список литературы", level=1)
    
    # Добавляем источники в неправильном порядке (Петров перед Ивановым)
    add_reference_entry(doc, "2. Петров, А. Б. Книга о психологии. М., 2020.")
    add_reference_entry(doc, "1. Иванов, И. И. Учебник по психологии. М., 2019.")
    
    doc.add_heading("Заключение", level=1)
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_L_4_cyrillic_before_latin_docx(tmp_path_factory):
    """Латиница перед кириллицей в списке литературы (Л-4)."""
    path = tmp_path_factory.mktemp("fix") / "wrong_L_4_cyrillic_latin.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    doc.add_heading("Список литературы", level=1)
    
    # Сначала иностранный, потом русский — нарушение
    add_reference_entry(doc, "1. Smith, J. Psychology book. London, 2020.")
    add_reference_entry(doc, "2. Иванов, И. И. Учебник по психологии. М., 2019.")
    
    doc.add_heading("Заключение", level=1)
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_L_5_numbering_docx(tmp_path_factory):
    """Нарушение сплошной нумерации источников (Л-5)."""
    path = tmp_path_factory.mktemp("fix") / "wrong_L_5_numbering.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    doc.add_heading("Список литературы", level=1)
    
    # Пропуск номера: 1, 3 вместо 1, 2
    add_reference_entry(doc, "1. Иванов, И. И. Учебник. М., 2019.")
    add_reference_entry(doc, "3. Петров, П. П. Книга. М., 2020.")
    
    doc.add_heading("Заключение", level=1)
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_L_8_old_sources_docx(tmp_path_factory):
    """Источники старше 10 лет (Л-8)."""
    path = tmp_path_factory.mktemp("fix") / "wrong_L_8_old_sources.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    doc.add_heading("Список литературы", level=1)
    
    # Все источники старые (больше 10 лет)
    for i in range(1, 11):
        add_reference_entry(doc, f"{i}. Автор, А. А. Книга. М., 2010.")
    
    doc.add_heading("Заключение", level=1)
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_L_9_author_format_docx(tmp_path_factory):
    """Неправильный формат автора (Л-9)."""
    path = tmp_path_factory.mktemp("fix") / "wrong_L_9_author_format.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    doc.add_heading("Список литературы", level=1)
    
    # Неправильный формат: без запятой или без пробелов
    add_reference_entry(doc, "1. Иванов И.И. Учебник. М., 2019.")
    
    doc.add_heading("Заключение", level=1)
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_L_10_url_no_date_docx(tmp_path_factory):
    """URL без даты обращения (Л-10)."""
    path = tmp_path_factory.mktemp("fix") / "wrong_L_10_url_no_date.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    doc.add_heading("Список литературы", level=1)
    
    # URL без даты обращения
    add_reference_entry(doc, "1. Сайт [Электронный ресурс]. URL: https://example.com")
    
    doc.add_heading("Заключение", level=1)
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_L_11_invalid_reference_docx(tmp_path_factory):
    """Ссылка на несуществующий источник (Л-11)."""
    path = tmp_path_factory.mktemp("fix") / "wrong_L_11_invalid_ref.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    
    # Ссылка на источник 99, которого нет в списке
    p = doc.add_paragraph()
    run = p.add_run("Текст со ссылкой [99].")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    set_paragraph_spacing(p, 420)
    set_first_line_indent(p, 720)
    set_alignment(p, 'both')
    
    doc.add_heading("Список литературы", level=1)
    # Только один источник
    add_reference_entry(doc, "1. Иванов, И. И. Учебник. М., 2019.")
    
    doc.add_heading("Заключение", level=1)
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_L_12_hyphen_instead_of_dash_docx(tmp_path_factory):
    """Дефис вместо длинного тире в библиографии (Л-12)."""
    path = tmp_path_factory.mktemp("fix") / "wrong_L_12_hyphen.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    doc.add_heading("Список литературы", level=1)
    
    # Используем дефис как разделитель (нарушение)
    add_reference_entry(doc, "1. Иванов, И. И. - М.: Издательство, 2019.")
    
    doc.add_heading("Заключение", level=1)
    doc.save(path)
    return path


# ── Фикстуры для тестов структуры (С-*) ──────────────────────────────────────

def add_page_break(doc):
    """Добавляет разрыв страницы."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore')
    pb.set(qn('w:val'), '1')
    pPr.append(pb)


@pytest.fixture(scope="session")
def wrong_C_2_appendix_no_reference_docx(tmp_path_factory):
    """Нарушение С-2: есть приложение, но нет ссылки на него в тексте."""
    path = tmp_path_factory.mktemp("fix") / "wrong_C_2_appendix_no_ref.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    add_correct_paragraph(doc, "Основной текст без ссылок на приложения.")
    
    doc.add_heading("Глава 1. Теоретические основы", level=1)
    add_correct_paragraph(doc, "Текст главы.")
    
    doc.add_heading("Заключение", level=1)
    add_correct_paragraph(doc, "Выводы.")
    
    doc.add_heading("Список литературы", level=1)
    
    # Добавляем приложение без ссылки в тексте
    doc.add_heading("Приложение А", level=1)
    add_correct_paragraph(doc, "Дополнительные материалы.")
    
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_C_3_section_no_page_break_docx(tmp_path_factory):
    """Нарушение С-3: раздел начинается без разрыва страницы."""
    path = tmp_path_factory.mktemp("fix") / "wrong_C_3_no_page_break.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    add_correct_paragraph(doc, "Текст введения.")
    
    # Глава без разрыва страницы перед ней (нарушение С-3)
    doc.add_heading("Глава 1. Теоретические основы", level=1)
    add_correct_paragraph(doc, "Текст главы.")
    
    doc.add_heading("Заключение", level=1)
    add_correct_paragraph(doc, "Выводы.")
    
    doc.add_heading("Список литературы", level=1)
    
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_C_4_paragraph_with_page_break_docx(tmp_path_factory):
    """Нарушение С-4: параграф начинается с новой страницы."""
    path = tmp_path_factory.mktemp("fix") / "wrong_C_4_para_page_break.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    add_correct_paragraph(doc, "Текст введения.")
    
    doc.add_heading("Глава 1. Теоретические основы", level=1)
    add_correct_paragraph(doc, "Текст главы.")
    
    # Параграф с разрывом страницы (нарушение С-4)
    # Добавляем pageBreakBefore непосредственно к заголовку параграфа
    p = doc.add_heading("1.1. Первый параграф", level=2)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore')
    pb.set(qn('w:val'), '1')
    pPr.append(pb)
    
    add_correct_paragraph(doc, "Текст параграфа.")
    
    doc.add_heading("Заключение", level=1)
    add_correct_paragraph(doc, "Выводы.")
    
    doc.add_heading("Список литературы", level=1)
    
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_C_6_paragraph_numbering_docx(tmp_path_factory):
    """Нарушение С-6: неправильная нумерация параграфа."""
    path = tmp_path_factory.mktemp("fix") / "wrong_C_6_numbering.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    add_correct_paragraph(doc, "Текст введения.")
    
    doc.add_heading("Глава 1. Теоретические основы", level=1)
    
    # Параграф без правильной нумерации (нарушение С-6)
    doc.add_heading("Первый параграф", level=2)
    add_correct_paragraph(doc, "Текст параграфа без нумерации.")
    
    doc.add_heading("Заключение", level=1)
    add_correct_paragraph(doc, "Выводы.")
    
    doc.add_heading("Список литературы", level=1)
    
    doc.save(path)
    return path


@pytest.fixture(scope="session")
def wrong_C_10_subheading_in_paragraph_docx(tmp_path_factory):
    """Нарушение С-10: подзаголовок внутри параграфа."""
    path = tmp_path_factory.mktemp("fix") / "wrong_C_10_subheading.docx"
    doc = make_base_doc()
    doc.add_heading("Введение", level=1)
    add_correct_paragraph(doc, "Текст введения.")
    
    doc.add_heading("Глава 1. Теоретические основы", level=1)
    doc.add_heading("1.1. Первый параграф", level=2)
    add_correct_paragraph(doc, "Начало текста параграфа.")
    
    # Подзаголовок внутри параграфа (нарушение С-10)
    doc.add_heading("1.1.1. Подраздел", level=3)
    add_correct_paragraph(doc, "Текст подраздела.")
    
    doc.add_heading("Заключение", level=1)
    add_correct_paragraph(doc, "Выводы.")
    
    doc.add_heading("Список литературы", level=1)
    
    doc.save(path)
    return path
