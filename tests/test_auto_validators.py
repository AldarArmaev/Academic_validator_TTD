# tests/test_auto_validators.py
"""
Автотесты для Academic Validator.

Каждый тест сам создаёт .docx-файл в tmp_path (pytest-фикстура),
запускает validate_format() и проверяет наличие / отсутствие нужного кода ошибки.

Запуск:
    pytest tests/test_auto_validators.py -v
"""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.shared import Twips

from src.validators.format_validator import validate_format

# ---------------------------------------------------------------------------
# Правила (копия из conftest / основного кода)
# ---------------------------------------------------------------------------

RULES: dict = {
    "font": {"family": "Times New Roman", "size_pt": 14},
    "paragraph": {"line_spacing_twips": 420},
    "margins_dxa": {"left": 1701, "right": 567, "top": 1134, "bottom": 1134},
    "tolerances": {"pt": 0.5, "dxa": 20},
    "required_sections": ["введение", "заключение", "список литературы"],
    "chapter_heading_pattern": r"^Глава \d+\.\s.+",
    "paragraph_heading_pattern": r"^\d+\.\d+(\.\d+)?\s.+",
    "references": {"min_sources": 40},
    "volume": {
        "total_chars_min": 90_000,
        "total_chars_max": 108_000,
        "theory_chapter_chars_min": 27_000,
        "theory_chapter_chars_max": 36_000,
        "empirical_chapter_chars_min": 45_000,
        "empirical_chapter_chars_max": 54_000,
    },
}

# ---------------------------------------------------------------------------
# Helpers для построения документов
# ---------------------------------------------------------------------------


def _save(doc: Document, path: Path) -> str:
    doc.save(str(path))
    return str(path)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Добавляет заголовок и выравнивает по центру."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_spacing(para, line_twips: int) -> None:
    """Устанавливает межстрочный интервал в twips."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(line_twips))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _set_alignment(para, value: str) -> None:
    """value: 'both' | 'left' | 'center' | 'right'"""
    pPr = para._p.get_or_add_pPr()
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), value)
    pPr.append(jc)


def _set_first_line_indent(para, dxa: int) -> None:
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), str(dxa))
    pPr.append(ind)


def _set_space_before_after(para, before: int, after: int) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


def _set_margins(doc: Document, left_cm, right_cm, top_cm, bottom_cm) -> None:
    section = doc.sections[0]
    section.left_margin = Cm(left_cm)
    section.right_margin = Cm(right_cm)
    section.top_margin = Cm(top_cm)
    section.bottom_margin = Cm(bottom_cm)


def _correct_margins(doc: Document) -> None:
    """Поля по методичке: левое 3 см, правое 1 см, верхнее/нижнее 2 см."""
    _set_margins(doc, 3.0, 1.0, 2.0, 2.0)


def _add_body_para(
    doc: Document,
    text: str,
    font_name: str = "Times New Roman",
    font_size_pt: float = 14,
    spacing_twips: int = 420,
    alignment: str = "both",
    first_line_dxa: int = 720,
) -> None:
    """Добавляет обычный абзац с заданным форматированием."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    _set_spacing(para, spacing_twips)
    _set_alignment(para, alignment)
    _set_first_line_indent(para, first_line_dxa)


def _minimal_correct_doc(doc: Document | None = None) -> Document:
    """
    Создаёт минимально корректный документ (не нарушающий Ф-1..Ф-6,
    Ф-4, С-1).  Используется как базис для «позитивных» тестов.
    """
    if doc is None:
        doc = Document()
    _correct_margins(doc)

    for section_name in ["Содержание", "Введение", "Заключение", "Список литературы"]:
        _add_heading(doc, section_name, level=1)
        _add_body_para(doc, f"Текст раздела «{section_name}».")

    return doc


# ---------------------------------------------------------------------------
# Ф-1: шрифт основного текста
# ---------------------------------------------------------------------------


class TestF1Font:
    """Ф-1 — Times New Roman, 14 пт."""

    def test_wrong_font_name_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        para = doc.add_paragraph()
        run = para.add_run("Неправильный шрифт Arial")
        run.font.name = "Arial"
        run.font.size = Pt(14)
        _set_spacing(para, 420)
        _set_alignment(para, "both")
        _set_first_line_indent(para, 720)

        path = _save(doc, tmp_path / "f1_wrong_font.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-1" in codes, "Ошибка Ф-1 не обнаружена при неверном шрифте"

    def test_wrong_font_size_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        para = doc.add_paragraph()
        run = para.add_run("Размер шрифта 12 пт — нарушение")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        _set_spacing(para, 420)
        _set_alignment(para, "both")
        _set_first_line_indent(para, 720)

        path = _save(doc, tmp_path / "f1_wrong_size.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-1" in codes, "Ошибка Ф-1 не обнаружена при неверном размере шрифта"

    def test_correct_font_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "Корректный текст Times New Roman 14 пт.")
        path = _save(doc, tmp_path / "f1_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Ф-1" for e in report.errors), "Ложное Ф-1"


# ---------------------------------------------------------------------------
# Ф-2: межстрочный интервал
# ---------------------------------------------------------------------------


class TestF2Spacing:
    """Ф-2 — 1.5 (420 twips)."""

    def test_wrong_spacing_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        para = doc.add_paragraph()
        run = para.add_run("Одинарный интервал — нарушение")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        _set_spacing(para, 240)   # 1.0 вместо 1.5
        _set_alignment(para, "both")
        _set_first_line_indent(para, 720)

        path = _save(doc, tmp_path / "f2_wrong.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-2" in codes, "Ошибка Ф-2 не обнаружена"

    def test_double_spacing_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        para = doc.add_paragraph()
        run = para.add_run("Двойной интервал — нарушение")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        _set_spacing(para, 480)   # 2.0
        _set_alignment(para, "both")
        _set_first_line_indent(para, 720)

        path = _save(doc, tmp_path / "f2_double.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-2" in codes, "Ошибка Ф-2 не обнаружена при двойном интервале"

    def test_correct_spacing_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "Полуторный интервал — всё ок.", spacing_twips=420)
        path = _save(doc, tmp_path / "f2_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Ф-2" for e in report.errors), "Ложное Ф-2"


# ---------------------------------------------------------------------------
# Ф-3: выравнивание по ширине
# ---------------------------------------------------------------------------


class TestF3Alignment:
    """Ф-3 — выравнивание по ширине (both)."""

    def test_left_alignment_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        para = doc.add_paragraph()
        run = para.add_run("Выравнивание влево — нарушение")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        _set_spacing(para, 420)
        _set_alignment(para, "left")
        _set_first_line_indent(para, 720)

        path = _save(doc, tmp_path / "f3_left.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-3" in codes, "Ошибка Ф-3 не обнаружена"

    def test_center_alignment_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        para = doc.add_paragraph()
        run = para.add_run("Выравнивание по центру — нарушение для основного текста")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        _set_spacing(para, 420)
        _set_alignment(para, "center")
        _set_first_line_indent(para, 720)

        path = _save(doc, tmp_path / "f3_center.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-3" in codes, "Ошибка Ф-3 не обнаружена при центрировании"

    def test_correct_alignment_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "Выравнивание по ширине — корректно.", alignment="both")
        path = _save(doc, tmp_path / "f3_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Ф-3" for e in report.errors), "Ложное Ф-3"


# ---------------------------------------------------------------------------
# Ф-4: поля документа
# ---------------------------------------------------------------------------


class TestF4Margins:
    """Ф-4 — поля: левое 3 см, правое 1 см, верхнее/нижнее 2 см."""

    def test_wrong_left_margin_detected(self, tmp_path):
        doc = Document()
        _set_margins(doc, 2.0, 1.0, 2.0, 2.0)   # левое 2 вместо 3
        for s in ["Введение", "Заключение", "Список литературы"]:
            _add_heading(doc, s, level=1)
            _add_body_para(doc, "Текст.")

        path = _save(doc, tmp_path / "f4_left_margin.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-4" in codes, "Ошибка Ф-4 не обнаружена при неверном левом поле"

    def test_wrong_right_margin_detected(self, tmp_path):
        doc = Document()
        _set_margins(doc, 3.0, 2.0, 2.0, 2.0)   # правое 2 вместо 1
        for s in ["Введение", "Заключение", "Список литературы"]:
            _add_heading(doc, s, level=1)
            _add_body_para(doc, "Текст.")

        path = _save(doc, tmp_path / "f4_right_margin.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-4" in codes, "Ошибка Ф-4 не обнаружена при неверном правом поле"

    def test_correct_margins_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        path = _save(doc, tmp_path / "f4_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Ф-4" for e in report.errors), "Ложное Ф-4"


# ---------------------------------------------------------------------------
# Ф-5: отступ первой строки
# ---------------------------------------------------------------------------


class TestF5FirstLineIndent:
    """Ф-5 — отступ красной строки 1,25 см (720 DXA)."""

    def test_zero_indent_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        para = doc.add_paragraph()
        run = para.add_run("Нет красной строки — нарушение")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        _set_spacing(para, 420)
        _set_alignment(para, "both")
        _set_first_line_indent(para, 0)   # отступ 0 — нарушение

        path = _save(doc, tmp_path / "f5_zero_indent.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-5" in codes, "Ошибка Ф-5 не обнаружена при нулевом отступе"

    def test_tab_indent_detected(self, tmp_path):
        """Отступ через Tab (большой — 1440 DXA) тоже нарушение."""
        doc = _minimal_correct_doc()
        para = doc.add_paragraph()
        run = para.add_run("Большой отступ через Tab")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        _set_spacing(para, 420)
        _set_alignment(para, "both")
        _set_first_line_indent(para, 1440)

        path = _save(doc, tmp_path / "f5_big_indent.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-5" in codes, "Ошибка Ф-5 не обнаружена при слишком большом отступе"

    def test_correct_indent_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "Корректный отступ 1,25 см.", first_line_dxa=720)
        path = _save(doc, tmp_path / "f5_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Ф-5" for e in report.errors), "Ложное Ф-5"


# ---------------------------------------------------------------------------
# Ф-6: интервалы до/после абзаца
# ---------------------------------------------------------------------------


class TestF6SpaceBeforeAfter:
    """Ф-6 — интервалы до и после абзаца должны быть 0."""

    def test_space_before_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        para = doc.add_paragraph()
        run = para.add_run("Интервал до абзаца 12 пт — нарушение")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        _set_spacing(para, 420)
        _set_alignment(para, "both")
        _set_first_line_indent(para, 720)
        _set_space_before_after(para, before=240, after=0)

        path = _save(doc, tmp_path / "f6_space_before.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-6" in codes, "Ошибка Ф-6 не обнаружена при ненулевом интервале до"

    def test_space_after_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        para = doc.add_paragraph()
        run = para.add_run("Интервал после абзаца — нарушение")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        _set_spacing(para, 420)
        _set_alignment(para, "both")
        _set_first_line_indent(para, 720)
        _set_space_before_after(para, before=0, after=200)

        path = _save(doc, tmp_path / "f6_space_after.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-6" in codes, "Ошибка Ф-6 не обнаружена при ненулевом интервале после"


# ---------------------------------------------------------------------------
# С-1: обязательные разделы
# ---------------------------------------------------------------------------


class TestC1RequiredSections:
    """С-1 — обязательные разделы: Введение, Заключение, Список литературы."""

    def test_missing_introduction_detected(self, tmp_path):
        doc = Document()
        _correct_margins(doc)
        _add_heading(doc, "Заключение", level=1)
        _add_body_para(doc, "Текст.")
        _add_heading(doc, "Список литературы", level=1)
        _add_body_para(doc, "Источник.")

        path = _save(doc, tmp_path / "c1_no_intro.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "С-1" in codes, "Ошибка С-1 не обнаружена при отсутствии Введения"

    def test_missing_conclusion_detected(self, tmp_path):
        doc = Document()
        _correct_margins(doc)
        _add_heading(doc, "Введение", level=1)
        _add_body_para(doc, "Текст.")
        _add_heading(doc, "Список литературы", level=1)
        _add_body_para(doc, "Источник.")

        path = _save(doc, tmp_path / "c1_no_conclusion.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "С-1" in codes, "Ошибка С-1 не обнаружена при отсутствии Заключения"

    def test_missing_references_detected(self, tmp_path):
        doc = Document()
        _correct_margins(doc)
        _add_heading(doc, "Введение", level=1)
        _add_body_para(doc, "Текст.")
        _add_heading(doc, "Заключение", level=1)
        _add_body_para(doc, "Текст.")

        path = _save(doc, tmp_path / "c1_no_refs.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "С-1" in codes, "Ошибка С-1 не обнаружена при отсутствии Списка литературы"

    def test_all_sections_present_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        path = _save(doc, tmp_path / "c1_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "С-1" for e in report.errors), "Ложное С-1"


# ---------------------------------------------------------------------------
# С-5: формат заголовков глав «Глава N. Название»
# ---------------------------------------------------------------------------


class TestC5ChapterHeadingFormat:
    """С-5 — заголовки глав должны соответствовать «Глава N. Название»."""

    def test_wrong_chapter_format_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "1. Теоретические основы", level=1)   # без слова «Глава»

        path = _save(doc, tmp_path / "c5_wrong_format.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "С-5" in codes, "Ошибка С-5 не обнаружена"

    def test_correct_chapter_format_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Глава 1. Теоретические основы", level=1)
        _add_body_para(doc, "Содержание главы.")

        path = _save(doc, tmp_path / "c5_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "С-5" for e in report.errors), "Ложное С-5"


# ---------------------------------------------------------------------------
# С-6: нумерация параграфов «1.1. Название»
# ---------------------------------------------------------------------------


class TestC6ParagraphNumbering:
    """С-6 — параграфы нумеруются «1.1.» или «1.1.1.»."""

    def test_unnumbered_paragraph_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Теоретический раздел", level=2)   # нет числа

        path = _save(doc, tmp_path / "c6_wrong.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "С-6" in codes, "Ошибка С-6 не обнаружена"

    def test_correct_paragraph_numbering_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "1.1 Теоретический обзор", level=2)
        _add_body_para(doc, "Содержание параграфа.")

        path = _save(doc, tmp_path / "c6_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "С-6" for e in report.errors), "Ложное С-6"


# ---------------------------------------------------------------------------
# С-7: заголовки не должны быть bold/italic/underline
# ---------------------------------------------------------------------------


class TestC7HeadingFormatting:
    """С-7 — заголовки без выделений."""

    def test_bold_heading_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        p = doc.add_heading("Глава 1. Теория", level=1)
        for run in p.runs:
            run.font.bold = True

        path = _save(doc, tmp_path / "c7_bold.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "С-7" in codes, "Ошибка С-7 не обнаружена при жирном заголовке"

    def test_italic_heading_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        p = doc.add_heading("Глава 2. Эмпирика", level=1)
        for run in p.runs:
            run.font.italic = True

        path = _save(doc, tmp_path / "c7_italic.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "С-7" in codes, "Ошибка С-7 не обнаружена при курсивном заголовке"

    def test_underline_heading_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        p = doc.add_heading("Заключение", level=1)
        for run in p.runs:
            run.font.underline = True

        path = _save(doc, tmp_path / "c7_underline.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "С-7" in codes, "Ошибка С-7 не обнаружена при подчёркнутом заголовке"


# ---------------------------------------------------------------------------
# С-9: нет точки в конце заголовка
# ---------------------------------------------------------------------------


class TestC9HeadingTrailingDot:
    """С-9 — точки в конце заголовков не ставятся."""

    def test_dot_at_end_of_heading_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Глава 1. Теоретические основы.", level=1)

        path = _save(doc, tmp_path / "c9_dot.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "С-9" in codes, "Ошибка С-9 не обнаружена при точке в конце заголовка"

    def test_no_dot_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Глава 1. Теоретические основы", level=1)
        path = _save(doc, tmp_path / "c9_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "С-9" for e in report.errors), "Ложное С-9"


# ---------------------------------------------------------------------------
# Т-1: подпись таблицы «Таблица N» над таблицей, по правому краю
# ---------------------------------------------------------------------------


class TestT1TableCaption:
    """Т-1 — подпись таблицы над таблицей, выравнивание по правому краю."""

    def test_table_without_caption_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        # Таблица без подписи
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"

        path = _save(doc, tmp_path / "t1_no_caption.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Т-1" in codes, "Ошибка Т-1 не обнаружена при отсутствии подписи таблицы"

    def test_caption_wrong_alignment_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        caption = doc.add_paragraph("Таблица 1")
        _set_alignment(caption, "left")   # должна быть right
        doc.add_table(rows=2, cols=2)

        path = _save(doc, tmp_path / "t1_wrong_align.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Т-1" in codes, "Ошибка Т-1 не обнаружена при неверном выравнивании подписи"


# ---------------------------------------------------------------------------
# Т-3: нет точки после номера и названия таблицы
# ---------------------------------------------------------------------------


class TestT3TableDot:
    """Т-3 — без точки после номера и названия таблицы."""

    def test_dot_at_end_of_table_caption_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        caption = doc.add_paragraph("Таблица 1.")   # точка после номера — нарушение
        _set_alignment(caption, "right")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "X"

        path = _save(doc, tmp_path / "t3_dot.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Т-1" in codes or "Т-3" in codes, "Ошибка Т-1/Т-3 не обнаружена"


# ---------------------------------------------------------------------------
# Т-4: шрифт в таблице Times New Roman 11–12 пт
# ---------------------------------------------------------------------------


class TestT4TableFont:
    """Т-4 — шрифт в таблице Times New Roman, 11–12 пт."""

    def test_wrong_font_in_table_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        caption = doc.add_paragraph("Таблица 1")
        _set_alignment(caption, "right")

        table = doc.add_table(rows=2, cols=2)
        run = table.cell(0, 0).paragraphs[0].add_run("Заголовок")
        run.font.name = "Arial"        # неверный шрифт
        run.font.size = Pt(12)

        path = _save(doc, tmp_path / "t4_wrong_font.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Т-4" in codes, "Ошибка Т-4 не обнаружена при неверном шрифте в таблице"

    def test_wrong_font_size_in_table_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        caption = doc.add_paragraph("Таблица 1")
        _set_alignment(caption, "right")

        table = doc.add_table(rows=2, cols=2)
        run = table.cell(0, 0).paragraphs[0].add_run("Текст")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)   # 14 пт вместо 11–12

        path = _save(doc, tmp_path / "t4_wrong_size.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Т-4" in codes, "Ошибка Т-4 не обнаружена при неверном размере шрифта в таблице"


# ---------------------------------------------------------------------------
# Т-6: сквозная нумерация
# ---------------------------------------------------------------------------


class TestT6SequentialNumbering:
    """Т-6 — сквозная нумерация таблиц и рисунков."""

    def test_broken_table_numbering_detected(self, tmp_path):
        doc = _minimal_correct_doc()

        for num in (1, 3):     # пропущена таблица 2
            cap = doc.add_paragraph(f"Таблица {num}")
            _set_alignment(cap, "right")
            t = doc.add_table(rows=2, cols=2)
            t.cell(0, 0).text = "x"

        path = _save(doc, tmp_path / "t6_broken.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Т-6" in codes, "Ошибка Т-6 не обнаружена при нарушении нумерации таблиц"


# ---------------------------------------------------------------------------
# Т-12: запятая в дробных числах
# ---------------------------------------------------------------------------


class TestT12DecimalSeparator:
    """Т-12 — десятичный разделитель — запятая, не точка."""

    def test_decimal_point_in_table_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        cap = doc.add_paragraph("Таблица 1")
        _set_alignment(cap, "right")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).paragraphs[0].add_run("3.14")   # точка вместо запятой

        path = _save(doc, tmp_path / "t12_decimal_point.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Т-12" in codes, "Ошибка Т-12 не обнаружена при точке как разделителе"

    def test_decimal_comma_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        cap = doc.add_paragraph("Таблица 1")
        _set_alignment(cap, "right")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).paragraphs[0].add_run("3,14")   # правильно

        path = _save(doc, tmp_path / "t12_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Т-12" for e in report.errors), "Ложное Т-12"


# ---------------------------------------------------------------------------
# Л-1: ссылки в квадратных скобках [N] или [N, с. X]
# ---------------------------------------------------------------------------


class TestL1CitationFormat:
    """Л-1 — ссылки в квадратных скобках."""

    def test_round_bracket_citation_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "Как указывает автор (5), проблема актуальна.")

        path = _save(doc, tmp_path / "l1_round.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Л-1" in codes, "Ошибка Л-1 не обнаружена при круглых скобках"

    def test_square_bracket_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "Как указывает автор [5], проблема актуальна.")

        path = _save(doc, tmp_path / "l1_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Л-1" for e in report.errors), "Ложное Л-1"


# ---------------------------------------------------------------------------
# Л-3: порядок множественных ссылок [4; 12; 25]
# ---------------------------------------------------------------------------


class TestL3MultiCitationOrder:
    """Л-3 — несколько источников перечисляются в возрастающем порядке."""

    def test_wrong_order_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "По данным исследований [12; 4; 25] можно утверждать...")

        path = _save(doc, tmp_path / "l3_wrong_order.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Л-3" in codes, "Ошибка Л-3 не обнаружена при неверном порядке ссылок"

    def test_correct_order_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "По данным исследований [4; 12; 25] можно утверждать...")

        path = _save(doc, tmp_path / "l3_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Л-3" for e in report.errors), "Ложное Л-3"


# ---------------------------------------------------------------------------
# Л-5: сплошная нумерация источников
# ---------------------------------------------------------------------------


class TestL5SequentialBibliography:
    """Л-5 — источники нумеруются сплошной нумерацией."""

    def test_broken_numbering_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Список литературы", level=1)
        doc.add_paragraph("1. Иванов, И. И. Психология. — М. : Наука, 2020. — 200 с.")
        doc.add_paragraph("3. Петров, П. П. Теория. — СПб. : Питер, 2021. — 150 с.")

        path = _save(doc, tmp_path / "l5_broken.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Л-5" in codes, "Ошибка Л-5 не обнаружена при нарушении нумерации"

    def test_correct_numbering_no_false_positive(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Список литературы", level=1)
        doc.add_paragraph("1. Иванов, И. И. Психология. — М. : Наука, 2020. — 200 с.")
        doc.add_paragraph("2. Петров, П. П. Теория. — СПб. : Питер, 2021. — 150 с.")

        path = _save(doc, tmp_path / "l5_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Л-5" for e in report.errors), "Ложное Л-5"


# ---------------------------------------------------------------------------
# Л-7: минимум 40 источников
# ---------------------------------------------------------------------------


class TestL7MinSources:
    """Л-7 — не менее 40 источников в списке литературы."""

    def test_too_few_sources_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Список литературы", level=1)
        for i in range(1, 11):    # только 10 источников
            doc.add_paragraph(
                f"{i}. Автор, А. А. Название. — М. : Наука, 2020. — {i * 10} с."
            )

        path = _save(doc, tmp_path / "l7_few.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Л-7" in codes, "Ошибка Л-7 не обнаружена при малом числе источников"

    def test_enough_sources_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Список литературы", level=1)
        for i in range(1, 42):    # 41 источник
            doc.add_paragraph(
                f"{i}. Автор, А. А. Название работы {i}. — М. : Наука, 2020. — 200 с."
            )

        path = _save(doc, tmp_path / "l7_enough.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Л-7" for e in report.errors), "Ложное Л-7"


# ---------------------------------------------------------------------------
# Л-9: формат автора «Фамилия, И. О.»
# ---------------------------------------------------------------------------


class TestL9AuthorFormat:
    """Л-9 — автор указывается в формате «Фамилия, И. О.»"""

    def test_wrong_author_format_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Список литературы", level=1)
        # Инициалы перед фамилией — нарушение
        doc.add_paragraph("1. И. И. Иванов. Психология. — М. : Наука, 2020. — 200 с.")

        path = _save(doc, tmp_path / "l9_wrong.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Л-9" in codes, "Ошибка Л-9 не обнаружена при неверном формате автора"

    def test_correct_author_format_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Список литературы", level=1)
        doc.add_paragraph(
            "1. Иванов, И. И. Психология личности. — М. : Наука, 2020. — 200 с."
        )

        path = _save(doc, tmp_path / "l9_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Л-9" for e in report.errors), "Ложное Л-9"


# ---------------------------------------------------------------------------
# Л-10: URL с датой обращения
# ---------------------------------------------------------------------------


class TestL10UrlAccessDate:
    """Л-10 — для URL указывается дата обращения."""

    def test_url_without_date_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Список литературы", level=1)
        doc.add_paragraph(
            "1. Сайт ИГУ. URL: https://isu.ru/ru/faculties/psy/"
        )

        path = _save(doc, tmp_path / "l10_no_date.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Л-10" in codes, "Ошибка Л-10 не обнаружена при URL без даты"

    def test_url_with_date_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Список литературы", level=1)
        doc.add_paragraph(
            "1. Сайт ИГУ. URL: https://isu.ru/ (дата обращения: 01.01.2024)"
        )

        path = _save(doc, tmp_path / "l10_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Л-10" for e in report.errors), "Ложное Л-10"


# ---------------------------------------------------------------------------
# Л-12: длинные тире в библиографии
# ---------------------------------------------------------------------------


class TestL12LongDash:
    """Л-12 — в библиографии используется тире «–», а не дефис «-»."""

    def test_hyphen_as_dash_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Список литературы", level=1)
        # Дефис вместо тире между числами
        doc.add_paragraph(
            "1. Иванов, И. И. Психология. - М. : Наука, 2020. - 200 с."
        )

        path = _save(doc, tmp_path / "l12_hyphen.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Л-12" in codes, "Ошибка Л-12 не обнаружена при дефисе вместо тире"

    def test_long_dash_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_heading(doc, "Список литературы", level=1)
        doc.add_paragraph(
            "1. Иванов, И. И. Психология личности. — М. : Наука, 2020. — 200 с."
        )

        path = _save(doc, tmp_path / "l12_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Л-12" for e in report.errors), "Ложное Л-12"


# ---------------------------------------------------------------------------
# Н-2: пробелы между инициалами
# ---------------------------------------------------------------------------


class TestN2InitialsSpacing:
    """Н-2 — между инициалами и перед фамилией ставится пробел."""

    def test_no_space_between_initials_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "Исследования А.Г.Асмолова подтверждают это.")

        path = _save(doc, tmp_path / "n2_no_space.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Н-2" in codes, "Ошибка Н-2 не обнаружена при отсутствии пробелов"

    def test_correct_initials_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "Исследования А. Г. Асмолова подтверждают это.")

        path = _save(doc, tmp_path / "n2_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Н-2" for e in report.errors), "Ложное Н-2"


# ---------------------------------------------------------------------------
# Н-4: угловые кавычки «»
# ---------------------------------------------------------------------------


class TestN4Quotes:
    """Н-4 — кавычки угловые «текст», а не "текст"."""

    def test_wrong_quotes_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, 'Автор использует термин "адаптация" в широком смысле.')

        path = _save(doc, tmp_path / "n4_wrong.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Н-4" in codes, "Ошибка Н-4 не обнаружена при неверных кавычках"

    def test_correct_quotes_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "Автор использует термин «адаптация» в широком смысле.")

        path = _save(doc, tmp_path / "n4_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Н-4" for e in report.errors), "Ложное Н-4"


# ---------------------------------------------------------------------------
# Н-5: тире между числами, а не дефис
# ---------------------------------------------------------------------------


class TestN5DashBetweenNumbers:
    """Н-5 — между числами/датами тире (–), а не дефис (-)."""

    def test_hyphen_between_years_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "В период 1990-2000 годов наблюдался рост.")

        path = _save(doc, tmp_path / "n5_hyphen.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Н-5" in codes, "Ошибка Н-5 не обнаружена при дефисе между годами"

    def test_dash_between_years_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "В период 1990–2000 годов наблюдался рост.")

        path = _save(doc, tmp_path / "n5_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Н-5" for e in report.errors), "Ложное Н-5"


# ---------------------------------------------------------------------------
# Н-6: аббревиатуры без расшифровки
# ---------------------------------------------------------------------------


class TestN6AbbreviationExplanation:
    """Н-6 — аббревиатура при первом употреблении расшифровывается."""

    def test_unexplained_abbreviation_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        # СПТ появляется без расшифровки
        _add_body_para(doc, "Использование СПТ в работе с группой эффективно.")

        path = _save(doc, tmp_path / "n6_no_explanation.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Н-6" in codes, "Ошибка Н-6 не обнаружена при нерасшифрованной аббревиатуре"

    def test_explained_abbreviation_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(
            doc,
            "Использование социально-психологического тренинга (СПТ) "
            "в работе с группой эффективно.",
        )

        path = _save(doc, tmp_path / "n6_explained.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Н-6" for e in report.errors), "Ложное Н-6"


# ---------------------------------------------------------------------------
# Л-2: повторная ссылка «там же»
# ---------------------------------------------------------------------------


class TestL2RepeatCitation:
    """Л-2 — повторная ссылка на тот же источник оформляется [там же, с. X]."""

    def test_double_ref_in_same_paragraph_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(
            doc,
            "Согласно теории [5, с. 10] … и далее [5, с. 25].",
        )

        path = _save(doc, tmp_path / "l2_double.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Л-2" in codes, "Ошибка Л-2 не обнаружена при двойной ссылке"

    def test_tam_zhe_no_error(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "Согласно теории [5, с. 10] … и далее [там же, с. 25].")

        path = _save(doc, tmp_path / "l2_correct.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Л-2" for e in report.errors), "Ложное Л-2"


# ---------------------------------------------------------------------------
# Со-1: содержание отражает все заголовки
# ---------------------------------------------------------------------------


class TestCo1TableOfContents:
    """Со-1 — содержание должно отражать все заголовки."""

    def test_missing_toc_detected(self, tmp_path):
        doc = Document()
        _correct_margins(doc)
        # Нет раздела «Содержание» / «Оглавление»
        _add_heading(doc, "Введение", level=1)
        _add_body_para(doc, "Текст.")
        _add_heading(doc, "Заключение", level=1)
        _add_body_para(doc, "Текст.")
        _add_heading(doc, "Список литературы", level=1)
        _add_body_para(doc, "Источник.")

        path = _save(doc, tmp_path / "co1_no_toc.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Со-1" in codes, "Ошибка Со-1 не обнаружена при отсутствии содержания"

    def test_toc_present_no_error(self, tmp_path):
        doc = Document()
        _correct_margins(doc)

        _add_heading(doc, "Содержание", level=1)
        doc.add_paragraph("Введение    3")
        doc.add_paragraph("Заключение    55")
        doc.add_paragraph("Список литературы    58")

        _add_heading(doc, "Введение", level=1)
        _add_body_para(doc, "Текст.")
        _add_heading(doc, "Заключение", level=1)
        _add_body_para(doc, "Текст.")
        _add_heading(doc, "Список литературы", level=1)
        _add_body_para(doc, "Источник.")

        path = _save(doc, tmp_path / "co1_present.docx")
        report = validate_format(path, RULES)
        assert all(e.code != "Со-1" for e in report.errors), "Ложное Со-1"


# ---------------------------------------------------------------------------
# П-1..П-3: приложения
# ---------------------------------------------------------------------------


def _add_page_break_before(para) -> None:
    """Добавляет w:pageBreakBefore к абзацу."""
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    pb.set(qn("w:val"), "1")
    pPr.append(pb)


class TestAppendixFormatting:
    """П-1..П-3 — оформление приложений."""

    def test_appendix_without_page_break_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "... (прил. 1)")
        # Добавляем приложение без разрыва страницы
        p = doc.add_paragraph("Приложение А")
        _set_alignment(p, "right")
        title_p = doc.add_paragraph("Название приложения")
        _set_alignment(title_p, "center")

        path = _save(doc, tmp_path / "p1_no_break.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "П-1" in codes, "Ошибка П-1 не обнаружена при отсутствии разрыва страницы"

    def test_appendix_wrong_alignment_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "... (прил. 1)")
        p = doc.add_paragraph("Приложение А")
        _add_page_break_before(p)
        _set_alignment(p, "center")   # должно быть right
        title_p = doc.add_paragraph("Название приложения")
        _set_alignment(title_p, "center")

        path = _save(doc, tmp_path / "p2_wrong_align.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "П-2" in codes, "Ошибка П-2 не обнаружена при неверном выравнивании приложения"

    def test_appendix_title_with_dot_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        _add_body_para(doc, "... (прил. 1)")
        p = doc.add_paragraph("Приложение А")
        _add_page_break_before(p)
        _set_alignment(p, "right")
        title_p = doc.add_paragraph("Название приложения.")   # точка — нарушение
        _set_alignment(title_p, "center")

        path = _save(doc, tmp_path / "p3_dot_in_title.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "П-3" in codes, "Ошибка П-3 не обнаружена при точке в названии приложения"


# ---------------------------------------------------------------------------
# Ф-11: общий объём работы
# ---------------------------------------------------------------------------


class TestF11Volume:
    """Ф-11 — объём 90 000–108 000 знаков с пробелами."""

    def test_too_short_detected(self, tmp_path):
        doc = _minimal_correct_doc()
        # Документ слишком короткий
        path = _save(doc, tmp_path / "f11_short.docx")
        report = validate_format(path, RULES)
        codes = [e.code for e in report.errors]
        assert "Ф-11" in codes, "Ошибка Ф-11 не обнаружена при слишком маленьком объёме"

    # def test_within_volume_no_error(self, tmp_path):
    #     doc = _minimal_correct_doc()
    #     # ~100 000 знаков: 1000 абзацев по 100 символов
    #     chunk = "А" * 95 + " "    # 96 символов с пробелом
    #     for _ in range(1050):
    #         _add_body_para(doc, chunk)
    #
    #     path = _save(doc, tmp_path / "f11_correct.docx")
    #     report = validate_format(path, RULES)
    #     assert all(e.code != "Ф-11" for e in report.errors), "Ложное Ф-11"

    # def test_too_long_detected(self, tmp_path):
    #     doc = _minimal_correct_doc()
    #     chunk = "Б" * 95 + " "
    #     for _ in range(1200):   # ~115 200 знаков — превышение
    #         _add_body_para(doc, chunk)
    #
    #     path = _save(doc, tmp_path / "f11_long.docx")
    #     report = validate_format(path, RULES)
    #     codes = [e.code for e in report.errors]
    #     assert "Ф-11" in codes, "Ошибка Ф-11 не обнаружена при превышении объёма"


# ---------------------------------------------------------------------------
# Итоговый «позитивный» тест: в корректном документе нет критических ошибок
# ---------------------------------------------------------------------------


class TestFullDocumentNoErrors:
    """Комплексный тест: полностью корректный документ не должен давать ошибок Ф-1..Ф-6."""

    def test_minimal_correct_doc_has_no_core_errors(self, tmp_path):
        """Ф-1, Ф-2, Ф-3, Ф-4, Ф-5, Ф-6 отсутствуют в корректном документе."""
        doc = _minimal_correct_doc()
        _add_body_para(
            doc,
            "Корректный абзац основного текста без нарушений форматирования.",
        )
        path = _save(doc, tmp_path / "full_correct.docx")
        report = validate_format(path, RULES)

        core_codes = {"Ф-1", "Ф-2", "Ф-3", "Ф-4", "Ф-5", "Ф-6"}
        found_core = {e.code for e in report.errors} & core_codes
        assert not found_core, (
            f"Найдены ложные ошибки: {found_core}\n"
            f"Все ошибки: {[e.code for e in report.errors]}"
        )