# src/validators/format_validator.py
"""Общий валидатор форматирования документов."""

from docx import Document
from docx.oxml.ns import qn
import re
import uuid
from datetime import datetime, timezone, timedelta
from typing import Any

from src.schemas import ValidationReport, ReportSummary, ReportError, ErrorLocation
from src.validators.font_validator import check_font_formatting


# ---------------------------------------------------------------------------
# ИСПРАВЛЕНИЕ №1 (Ф-6): вспомогательная функция для получения w:spacing
# ---------------------------------------------------------------------------

def _get_merged_spacing(pPr) -> dict:
    """
    Возвращает словарь всех атрибутов w:spacing, объединяя несколько
    элементов w:spacing в одном w:pPr (python-docx иногда создаёт их два:
    один для line, другой для before/after).
    """
    if pPr is None:
        return {}
    attrs: dict = {}
    for spacing_el in pPr.findall(qn('w:spacing')):
        for k, v in spacing_el.attrib.items():
            attrs[k] = v
    return attrs


def check_paragraph_formatting(doc: Document, rules: dict[str, Any]) -> list[ReportError]:
    """
    Проверяет форматирование абзацев.

    А) Межстрочный интервал (Ф-2)
    Б) Выравнивание (Ф-3)
    В) Отступ первой строки (Ф-5)
    Г) Интервалы до/после (Ф-6)
    """
    errors: list[ReportError] = []

    expected_line_spacing = rules["paragraph"]["line_spacing_twips"]  # 420
    tolerance_dxa = rules["tolerances"]["dxa"]  # 20
    expected_first_line_indent = 720  # DXA
    expected_before_after = 0  # twips

    heading_styles = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6"}

    chapter_heading_pattern = rules.get("chapter_heading_pattern", r"^Глава \d+\.\s.+")
    paragraph_heading_pattern = rules.get("paragraph_heading_pattern", r"^\d+\.\d+(\.\d+)?\s.+")

    for para_index, para in enumerate(doc.paragraphs):
        if para.style and para.style.name in heading_styles:
            continue
        if not para.text.strip():
            continue

        text_stripped = para.text.strip()
        if (text_stripped.startswith('Таблица') or
                text_stripped.startswith('Рис.') or
                text_stripped.startswith('Рисунок')):
            continue

        pPr = para._p.pPr
        if pPr is not None:
            outline_lvl = pPr.find(qn('w:outlineLvl'))
            if outline_lvl is not None:
                continue

        if (re.match(chapter_heading_pattern, text_stripped) or
                re.match(paragraph_heading_pattern, text_stripped)):
            continue

        # --- А) Межстрочный интервал (Ф-2) ---
        # ИСПРАВЛЕНИЕ №1: используем _get_merged_spacing вместо одного .find
        spacing_attrs = _get_merged_spacing(pPr)
        line_key = str(qn('w:line'))
        if line_key in spacing_attrs:
            try:
                actual_spacing = int(spacing_attrs[line_key])
                if abs(actual_spacing - expected_line_spacing) > tolerance_dxa:
                    errors.append(ReportError(
                        id=f"Ф-2-{para_index}",
                        code="Ф-2",
                        type="formatting",
                        severity="error",
                        location=ErrorLocation(
                            paragraph_index=para_index,
                            structural_path=f"Абзац {para_index + 1}"
                        ),
                        fragment=para.text[:100],
                        rule="Межстрочный интервал должен быть 1.5 (420 twips)",
                        rule_citation="§4.2, с. 47",
                        found_value=str(actual_spacing),
                        expected_value=str(expected_line_spacing),
                        recommendation="Установите межстрочный интервал 1.5"
                    ))
            except ValueError:
                pass

        # --- Б) Выравнивание (Ф-3) ---
        if pPr is not None:
            jc_el = pPr.find(qn('w:jc'))
            if jc_el is not None:
                alignment = jc_el.get(qn('w:val'))
                if alignment != "both":
                    errors.append(ReportError(
                        id=f"Ф-3-{para_index}",
                        code="Ф-3",
                        type="formatting",
                        severity="error",
                        location=ErrorLocation(
                            paragraph_index=para_index,
                            structural_path=f"Абзац {para_index + 1}"
                        ),
                        fragment=para.text[:100],
                        rule="Текст должен быть выровнен по ширине",
                        rule_citation="§4.2, с. 47",
                        found_value=alignment or "не задано",
                        expected_value="both",
                        recommendation="Установите выравнивание по ширине"
                    ))

        # --- В) Отступ первой строки (Ф-5) ---
        if pPr is not None:
            ind_el = pPr.find(qn('w:ind'))
            if ind_el is not None:
                first_line = ind_el.get(qn('w:firstLine'))
                if first_line is not None:
                    try:
                        actual_first_line = int(first_line)
                        if abs(actual_first_line - expected_first_line_indent) > tolerance_dxa:
                            errors.append(ReportError(
                                id=f"Ф-5-{para_index}",
                                code="Ф-5",
                                type="formatting",
                                severity="error",
                                location=ErrorLocation(
                                    paragraph_index=para_index,
                                    structural_path=f"Абзац {para_index + 1}"
                                ),
                                fragment=para.text[:100],
                                rule="Отступ первой строки должен быть 1.25 см (720 DXA)",
                                rule_citation="§4.2, с. 47",
                                found_value=str(actual_first_line),
                                expected_value=str(expected_first_line_indent),
                                recommendation="Установите отступ первой строки 1.25 см"
                            ))
                    except ValueError:
                        pass

        # --- Г) Интервалы до/после (Ф-6) ---
        # ИСПРАВЛЕНИЕ №1: читаем из объединённых атрибутов
        before_key = str(qn('w:before'))
        after_key = str(qn('w:after'))

        if before_key in spacing_attrs:
            try:
                actual_before = int(spacing_attrs[before_key])
                if actual_before != expected_before_after:
                    errors.append(ReportError(
                        id=f"Ф-6-before-{para_index}",
                        code="Ф-6",
                        type="formatting",
                        severity="error",
                        location=ErrorLocation(
                            paragraph_index=para_index,
                            structural_path=f"Абзац {para_index + 1}"
                        ),
                        fragment=para.text[:100],
                        rule="Интервалы до и после абзаца должны быть 0",
                        rule_citation="§4.2, с. 47",
                        found_value=str(actual_before),
                        expected_value=str(expected_before_after),
                        recommendation="Установите интервал перед абзацем 0"
                    ))
            except ValueError:
                pass

        if after_key in spacing_attrs:
            try:
                actual_after = int(spacing_attrs[after_key])
                if actual_after != expected_before_after:
                    errors.append(ReportError(
                        id=f"Ф-6-after-{para_index}",
                        code="Ф-6",
                        type="formatting",
                        severity="error",
                        location=ErrorLocation(
                            paragraph_index=para_index,
                            structural_path=f"Абзац {para_index + 1}"
                        ),
                        fragment=para.text[:100],
                        rule="Интервалы до и после абзаца должны быть 0",
                        rule_citation="§4.2, с. 47",
                        found_value=str(actual_after),
                        expected_value=str(expected_before_after),
                        recommendation="Установите интервал после абзаца 0"
                    ))
            except ValueError:
                pass

    return errors


def check_margins(doc: Document, rules: dict[str, Any]) -> list[ReportError]:
    errors: list[ReportError] = []

    section = doc.sections[0]
    EMU_PER_DXA = 635
    tolerance_dxa = rules["tolerances"]["dxa"]

    margins_config = {
        "left": section.left_margin,
        "right": section.right_margin,
        "top": section.top_margin,
        "bottom": section.bottom_margin
    }

    for margin_name, margin_emu in margins_config.items():
        expected_dxa = rules["margins_dxa"][margin_name]
        actual_dxa = round(margin_emu / EMU_PER_DXA)

        if abs(actual_dxa - expected_dxa) > tolerance_dxa:
            errors.append(ReportError(
                id=f"Ф-4-{margin_name}",
                code="Ф-4",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=0,
                    structural_path="Поля документа"
                ),
                fragment=f"Поле {margin_name}",
                rule=f"Поле {margin_name} должно быть {expected_dxa} DXA",
                rule_citation="§4.2, с. 47",
                found_value=str(actual_dxa),
                expected_value=str(expected_dxa),
                recommendation=f"Установите поле {margin_name} равным {expected_dxa} DXA"
            ))

    return errors


def validate_structure(doc: Document, rules: dict[str, Any]) -> list[ReportError]:
    errors: list[ReportError] = []

    def _get_effective_alignment(para, doc) -> str | None:
        pPr = para._p.find(qn('w:pPr'))
        jc_el = pPr.find(qn('w:jc')) if pPr is not None else None
        if jc_el is not None:
            return jc_el.get(qn('w:val'))
        if para.style and para.style.paragraph_format:
            fmt = para.style.paragraph_format
            if fmt.alignment is not None:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if fmt.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return "center"
                elif fmt.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                    return "left"
                elif fmt.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    return "both"
                elif fmt.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    return "right"
        if para.style and "Heading" in para.style.name:
            return "center"
        return None

    titles_lower = [
        p.text.strip().lower()
        for p in doc.paragraphs
        if p.style and p.style.name in ("Heading 1", "Heading 2")
    ]

    # С-1: Обязательные разделы
    for section_name in rules["required_sections"]:
        if not any(section_name.lower() in title for title in titles_lower):
            errors.append(ReportError(
                id=f"С-1-{section_name}",
                code="С-1",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=0,
                    structural_path="Структура документа"
                ),
                fragment=section_name,
                rule=f"Документ должен содержать раздел '{section_name}'",
                rule_citation="§3.2, с. 42",
                found_value="раздел отсутствует",
                expected_value=section_name,
                recommendation=f"Добавьте раздел '{section_name}' в документ"
            ))

    service_titles = ["введение", "заключение", "список литературы", "содержание"]

    pattern = rules.get("chapter_heading_pattern", r"^Глава \d+\.\s.+")
    paragraph_heading_pattern = rules.get("paragraph_heading_pattern", r"^\d+\.\d+(\.\d+)?\s.+")

    has_appendix = False
    appendix_ref_pattern = re.compile(r"\(прил\.\s*\d+\)", re.IGNORECASE)

    in_chapter = False
    chapter_start_idx = None

    for para_idx, para in enumerate(doc.paragraphs):
        if not para.style or "Heading" not in para.style.name:
            continue

        title = para.text.strip()
        title_lower = title.lower()

        if "приложен" in title_lower:
            has_appendix = True

        is_new_section = False
        for sec_name in service_titles:
            if sec_name in title_lower:
                is_new_section = True
                break

        if para.style.name == "Heading 1" and not any(s in title_lower for s in service_titles):
            is_new_section = True

        if is_new_section:
            pPr = para._p.find(qn('w:pPr'))
            has_page_break = False

            if pPr is not None:
                page_break_before = pPr.find(qn('w:pageBreakBefore'))
                if page_break_before is not None:
                    val = page_break_before.get(qn('w:val'))
                    if val in ('1', 'true', 'on'):
                        has_page_break = True

                if para_idx > 0:
                    prev_para = doc.paragraphs[para_idx - 1]
                    prev_pPr = prev_para._p.find(qn('w:pPr'))
                    if prev_pPr is not None:
                        for br in prev_pPr.findall(qn('w:br')):
                            br_type = br.get(qn('w:type'))
                            if br_type == 'page':
                                has_page_break = True
                                break

            if not has_page_break:
                errors.append(ReportError(
                    id=f"С-3-{para_idx}",
                    code="С-3",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path=f"Заголовок '{title[:50]}'"
                    ),
                    fragment=title[:100],
                    rule="Каждый новый раздел должен начинаться с новой страницы",
                    rule_citation="§4.2, с. 47",
                    found_value="нет разрыва страницы",
                    expected_value="разрыв страницы перед разделом",
                    recommendation="Добавьте разрыв страницы перед началом раздела"
                ))

        if para.style.name in ("Heading 2", "Heading 3"):
            pPr = para._p.find(qn('w:pPr'))
            has_page_break = False

            if pPr is not None:
                page_break_before = pPr.find(qn('w:pageBreakBefore'))
                if page_break_before is not None:
                    val = page_break_before.get(qn('w:val'))
                    if val is None or val in ('1', 'true', 'on'):
                        has_page_break = True

            if not has_page_break:
                xml_str = str(para._p.xml)
                if '<w:lastRenderedPageBreak' in xml_str:
                    has_page_break = True

            if has_page_break:
                errors.append(ReportError(
                    id=f"С-4-{para_idx}",
                    code="С-4",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path=f"Параграф '{title[:50]}'"
                    ),
                    fragment=title[:100],
                    rule="Параграфы не должны начинаться с новой страницы",
                    rule_citation="§4.2, с. 47",
                    found_value="есть разрыв страницы",
                    expected_value="нет разрыва страницы",
                    recommendation="Удалите разрыв страницы перед параграфом"
                ))

        if para.style.name == "Heading 1":
            if not any(s in title_lower for s in service_titles):
                if not re.match(pattern, title):
                    errors.append(ReportError(
                        id=f"С-5-{para_idx}",
                        code="С-5",
                        type="formatting",
                        severity="error",
                        location=ErrorLocation(
                            paragraph_index=para_idx,
                            structural_path=f"Заголовок {para_idx + 1}"
                        ),
                        fragment=title[:100],
                        rule="Заголовок главы должен соответствовать паттерну 'Глава N. Название'",
                        rule_citation="§3.3, с. 43",
                        found_value=title[:100],
                        expected_value="Глава N. Название",
                        recommendation="Измените формат заголовка главы"
                    ))

        if para.style.name in ("Heading 2", "Heading 3"):
            if not re.match(paragraph_heading_pattern, title):
                errors.append(ReportError(
                    id=f"С-6-{para_idx}",
                    code="С-6",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path=f"Параграф {para_idx + 1}"
                    ),
                    fragment=title[:100],
                    rule="Параграфы должны иметь нумерацию вида '1.1.' или '1.1.1.'",
                    rule_citation="§4.2, с. 47",
                    found_value=title[:100],
                    expected_value="N.N. Название или N.N.N. Название",
                    recommendation="Измените формат заголовка параграфа"
                ))

        has_formatting = False
        for run in para.runs:
            if run.font.bold or run.font.italic or run.font.underline:
                has_formatting = True
                break

        if has_formatting:
            errors.append(ReportError(
                id=f"С-7-{para_idx}",
                code="С-7",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=para_idx,
                    structural_path=f"Заголовок {para_idx + 1}"
                ),
                fragment=title[:100],
                rule="Заголовки не должны быть жирными, курсивом или подчёркнутыми",
                rule_citation="§3.3, с. 43",
                found_value="bold/italic/underline",
                expected_value="обычный текст",
                recommendation="Уберите жирность, курсив и подчёркивание из заголовка"
            ))

        effective_alignment = _get_effective_alignment(para, doc)

        if effective_alignment != "center":
            errors.append(ReportError(
                id=f"С-8-{para_idx}",
                code="С-8",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=para_idx,
                    structural_path=f"Заголовок {para_idx + 1}"
                ),
                fragment=title[:100],
                rule="Заголовки должны быть выровнены по центру",
                rule_citation="§3.3, с. 43",
                found_value=effective_alignment or "не задано",
                expected_value="center",
                recommendation="Установите выравнивание заголовка по центру"
            ))

        if title.endswith('.'):
            errors.append(ReportError(
                id=f"С-9-{para_idx}",
                code="С-9",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=para_idx,
                    structural_path=f"Заголовок {para_idx + 1}"
                ),
                fragment=title[:100],
                rule="В конце заголовка не должно быть точки",
                rule_citation="§3.3, с. 43",
                found_value=title[-10:] if len(title) > 10 else title,
                expected_value="без точки",
                recommendation="Удалите точку в конце заголовка"
            ))

    in_chapter = False
    chapter_start_idx = None

    for para_idx, para in enumerate(doc.paragraphs):
        if not para.style:
            continue

        style_name = para.style.name
        title = para.text.strip()

        if style_name == "Heading 1":
            in_chapter = True
            chapter_start_idx = para_idx
            continue

        if in_chapter and style_name == "Heading 2":
            if not re.match(r'^\d+\.\d+', title):
                errors.append(ReportError(
                    id=f"С-10-{para_idx}",
                    code="С-10",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path=f"Подзаголовок внутри главы {chapter_start_idx}"
                    ),
                    fragment=title[:100],
                    rule="Внутри параграфов не должно быть подзаголовков без правильной нумерации",
                    rule_citation="§4.2, с. 47",
                    found_value=title[:100],
                    expected_value="нумерация вида N.N или обычный текст",
                    recommendation="Удалите подзаголовок или добавьте правильную нумерацию"
                ))

        if in_chapter and style_name in ("Heading 3", "Heading 4", "Heading 5", "Heading 6"):
            if title:
                errors.append(ReportError(
                    id=f"С-10-sub-{para_idx}",
                    code="С-10",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path=f"Подзаголовок внутри главы {chapter_start_idx}"
                    ),
                    fragment=title[:100],
                    rule="Внутри параграфов не должно быть подзаголовков",
                    rule_citation="§4.2, с. 47",
                    found_value=title[:100],
                    expected_value="обычный текст без подзаголовков",
                    recommendation="Удалите подзаголовок или оформите его как обычный текст"
                ))

        if style_name == "Heading 1":
            in_chapter = False

    if has_appendix:
        full_text = "\n".join([p.text for p in doc.paragraphs])
        appendix_refs_found = bool(appendix_ref_pattern.search(full_text))

        if not appendix_refs_found:
            errors.append(ReportError(
                id="С-2-appx-ref",
                code="С-2",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=0,
                    structural_path="Приложения"
                ),
                fragment="Приложение",
                rule="Приложения должны иметь ссылки из текста в формате '(прил. N)'",
                rule_citation="§3.8, с. 44",
                found_value="ссылки на приложения отсутствуют",
                expected_value="(прил. N)",
                recommendation="Добавьте ссылки на приложения в тексте"
            ))

    return errors


def get_effective_font_size(run, doc) -> float | None:
    from docx.oxml.ns import qn

    if run.font.size is not None:
        return run.font.size.pt

    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        sz_el = rPr.find(qn('w:sz'))
        if sz_el is not None:
            val = sz_el.get(qn('w:val'))
            if val is not None:
                try:
                    return int(val) / 2.0
                except ValueError:
                    pass

        szcs_el = rPr.find(qn('w:szCs'))
        if szcs_el is not None:
            val = szcs_el.get(qn('w:val'))
            if val is not None:
                try:
                    return int(val) / 2.0
                except ValueError:
                    pass

    para = run._parent
    while para is not None and not hasattr(para, 'style'):
        para = para._parent

    if para is not None and hasattr(para, 'style') and para.style is not None:
        try:
            style_font_size = para.style.font.size
            if style_font_size is not None:
                return style_font_size.pt
        except Exception:
            pass

    return 12.0


def validate_tables(doc: Document, rules: dict[str, Any]) -> list[ReportError]:
    """Проверяет таблицы и рисунки (Т-1..Т-12)."""
    errors: list[ReportError] = []

    body = doc.element.body
    elements_flow = []
    para_index = 0

    for i, child in enumerate(body):
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            elements_flow.append(('paragraph', para_index, child))
            para_index += 1
        elif tag == 'tbl':
            elements_flow.append(('table', len([e for e in elements_flow if e[0] == 'table']), child))

    table_caption_pattern = re.compile(r'^Таблица\s*\d+', re.IGNORECASE)

    for elem_type, elem_idx, element in elements_flow:
        if elem_type != 'table':
            continue

        table_index = elem_idx
        caption_found = False

        for j in range(elements_flow.index((elem_type, elem_idx, element)) - 1, -1, -1):
            prev_type, prev_idx, prev_element = elements_flow[j]
            if prev_type != 'paragraph':
                break

            para_text = ''
            for t in prev_element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text:
                    para_text += t.text
            para_text = para_text.strip()

            if table_caption_pattern.match(para_text):
                caption_found = True

                pPr = prev_element.find(qn('w:pPr'))
                jc_el = pPr.find(qn('w:jc')) if pPr is not None else None
                alignment = jc_el.get(qn('w:val')) if jc_el is not None else None

                if alignment != 'right':
                    errors.append(ReportError(
                        id=f"Т-1-caption-align-{prev_idx}",
                        code="Т-1",
                        type="formatting",
                        severity="error",
                        location=ErrorLocation(
                            paragraph_index=prev_idx,
                            structural_path=f"Подпись таблицы {table_index + 1}"
                        ),
                        fragment=para_text[:100],
                        rule="Подпись таблицы должна быть выровнена по правому краю",
                        rule_citation="§4.5, с. 51",
                        found_value=alignment or "не задано",
                        expected_value="right",
                        recommendation="Установите выравнивание подписи таблицы по правому краю"
                    ))

                if para_text.endswith('.'):
                    errors.append(ReportError(
                        id=f"Т-1-caption-dot-{prev_idx}",
                        code="Т-1",
                        type="formatting",
                        severity="error",
                        location=ErrorLocation(
                            paragraph_index=prev_idx,
                            structural_path=f"Подпись таблицы {table_index + 1}"
                        ),
                        fragment=para_text[:100],
                        rule="В конце подписи таблицы не должно быть точки",
                        rule_citation="§4.5, с. 51",
                        found_value=para_text[-10:] if len(para_text) > 10 else para_text,
                        expected_value="без точки",
                        recommendation="Удалите точку в конце подписи таблицы"
                    ))

                break

            if para_text:
                break

        if not caption_found:
            errors.append(ReportError(
                id=f"Т-1-no-caption-{table_index}",
                code="Т-1",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=0,
                    structural_path=f"Таблица {table_index + 1}"
                ),
                fragment=f"Таблица {table_index + 1}",
                rule="Над таблицей должна быть подпись «Таблица N»",
                rule_citation="§4.5, с. 51",
                found_value="подпись отсутствует или расположена неверно",
                expected_value="Таблица N над таблицей",
                recommendation="Добавьте подпись «Таблица N» непосредственно над таблицей"
            ))

    table_title_below_pattern = re.compile(r'^Таблица\s*\d+\s*—\s*.+', re.IGNORECASE)

    for elem_type, elem_idx, element in elements_flow:
        if elem_type != 'table':
            continue

        table_index = elem_idx
        current_idx_in_flow = elements_flow.index((elem_type, elem_idx, element))

        for j in range(current_idx_in_flow + 1, len(elements_flow)):
            next_type, next_idx, next_element = elements_flow[j]

            if next_type != 'paragraph':
                break

            para_text = ''
            for t in next_element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text:
                    para_text += t.text
            para_text = para_text.strip()

            if not para_text:
                continue

            if table_title_below_pattern.match(para_text):
                pPr = next_element.find(qn('w:pPr'))
                jc_el = pPr.find(qn('w:jc')) if pPr is not None else None
                alignment = jc_el.get(qn('w:val')) if jc_el is not None else None

                if alignment != 'center':
                    errors.append(ReportError(
                        id=f"Т-2-align-{next_idx}",
                        code="Т-2",
                        type="formatting",
                        severity="error",
                        location=ErrorLocation(
                            paragraph_index=next_idx,
                            structural_path=f"Название таблицы {table_index + 1}"
                        ),
                        fragment=para_text[:100],
                        rule="Название таблицы должно быть выровнено по центру",
                        rule_citation="§4.4, с. 50-52",
                        found_value=alignment or "не задано",
                        expected_value="center",
                        recommendation="Установите выравнивание по центру"
                    ))
            break

    # Т-3
    for elem_type, elem_idx, element in elements_flow:
        if elem_type != 'table':
            continue

        table_index = elem_idx
        current_idx_in_flow = elements_flow.index((elem_type, elem_idx, element))

        for j in range(current_idx_in_flow + 1, len(elements_flow)):
            next_type, next_idx, next_element = elements_flow[j]

            if next_type != 'paragraph':
                break

            para_text = ''
            for t in next_element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text:
                    para_text += t.text
            para_text = para_text.strip()

            if not para_text:
                continue

            if table_title_below_pattern.match(para_text):
                dot_after_number = re.search(r'^Таблица\s*\d+\.', para_text, re.IGNORECASE)
                if dot_after_number:
                    errors.append(ReportError(
                        id=f"Т-3-dot-after-number-{next_idx}",
                        code="Т-3",
                        type="formatting",
                        severity="error",
                        location=ErrorLocation(
                            paragraph_index=next_idx,
                            structural_path=f"Название таблицы {table_index + 1}"
                        ),
                        fragment=para_text[:100],
                        rule="После номера таблицы не должна ставиться точка",
                        rule_citation="§4.4, с. 50-52",
                        found_value=para_text[:30],
                        expected_value="Таблица N (без точки)",
                        recommendation="Удалите точку после номера таблицы"
                    ))

                if para_text.rstrip().endswith('.'):
                    errors.append(ReportError(
                        id=f"Т-3-dot-at-end-{next_idx}",
                        code="Т-3",
                        type="formatting",
                        severity="error",
                        location=ErrorLocation(
                            paragraph_index=next_idx,
                            structural_path=f"Название таблицы {table_index + 1}"
                        ),
                        fragment=para_text[:100],
                        rule="В конце названия таблицы не должна ставиться точка",
                        rule_citation="§4.4, с. 50-52",
                        found_value=para_text[-20:] if len(para_text) > 20 else para_text,
                        expected_value="без точки в конце",
                        recommendation="Удалите точку в конце названия таблицы"
                    ))
            break

    # Т-4
    expected_font_name = "Times New Roman"
    min_font_size_pt = 11
    max_font_size_pt = 12

    tables_list = list(doc.tables)
    for table_idx, table in enumerate(tables_list):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        font_name = run.font.name

                        if font_name and font_name != expected_font_name:
                            errors.append(ReportError(
                                id=f"Т-4-font-name-{table_idx}-{row_idx}-{cell_idx}",
                                code="Т-4",
                                type="formatting",
                                severity="error",
                                location=ErrorLocation(
                                    paragraph_index=0,
                                    structural_path=f"Таблица {table_idx + 1}, ячейка [{row_idx + 1}, {cell_idx + 1}]"
                                ),
                                fragment=para.text[:50],
                                rule="Шрифт в таблице должен быть Times New Roman",
                                rule_citation="§4.5, с. 51",
                                found_value=font_name,
                                expected_value=expected_font_name,
                                recommendation="Установите шрифт Times New Roman"
                            ))

                        font_size = get_effective_font_size(run, doc)
                        if font_size is not None:
                            if font_size < min_font_size_pt or font_size > max_font_size_pt:
                                errors.append(ReportError(
                                    id=f"Т-4-font-size-{table_idx}-{row_idx}-{cell_idx}",
                                    code="Т-4",
                                    type="formatting",
                                    severity="error",
                                    location=ErrorLocation(
                                        paragraph_index=0,
                                        structural_path=f"Таблица {table_idx + 1}, ячейка [{row_idx + 1}, {cell_idx + 1}]"
                                    ),
                                    fragment=para.text[:50],
                                    rule="Размер шрифта в таблице должен быть 11-12 пт",
                                    rule_citation="§4.5, с. 51",
                                    found_value=str(font_size),
                                    expected_value=f"{min_font_size_pt}-{max_font_size_pt}",
                                    recommendation="Установите размер шрифта 11-12 пт"
                                ))

    # Т-5
    for table_idx, table in enumerate(tables_list):
        tbl_element = table._tbl
        tblPr = tbl_element.tblPr

        if tblPr is not None:
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None:
                w_type = tblW.get(qn('w:type'))
                if w_type not in ['auto', 'pct', 'nil']:
                    errors.append(ReportError(
                        id=f"Т-5-width-{table_idx}",
                        code="Т-5",
                        type="formatting",
                        severity="error",
                        location=ErrorLocation(
                            paragraph_index=0,
                            structural_path=f"Таблица {table_idx + 1}"
                        ),
                        fragment=f"Таблица {table_idx + 1}",
                        rule="Таблица должна иметь автоподбор по ширине окна",
                        rule_citation="§4.4, с. 50-52",
                        found_value=w_type or "не задано",
                        expected_value="auto или pct",
                        recommendation="Установите автоподбор ширины таблицы"
                    ))

        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    pPr = para._p.pPr
                    if pPr is not None:
                        jc_el = pPr.find(qn('w:jc'))
                        if jc_el is not None:
                            alignment = jc_el.get(qn('w:val'))
                            if alignment != 'center':
                                errors.append(ReportError(
                                    id=f"Т-5-cell-align-{table_idx}-{row_idx}-{cell_idx}",
                                    code="Т-5",
                                    type="formatting",
                                    severity="error",
                                    location=ErrorLocation(
                                        paragraph_index=0,
                                        structural_path=f"Таблица {table_idx + 1}, ячейка [{row_idx + 1}, {cell_idx + 1}]"
                                    ),
                                    fragment=para.text[:50],
                                    rule="Текст в ячейках таблицы должен быть выровнен по центру",
                                    rule_citation="§4.4, с. 50-52",
                                    found_value=alignment or "не задано",
                                    expected_value="center",
                                    recommendation="Установите выравнивание по центру в ячейке"
                                ))

    # Т-6: сквозная нумерация
    table_numbers = []
    figure_numbers = []

    for elem_type, elem_idx, element in elements_flow:
        if elem_type != 'table':
            continue

        current_idx_in_flow = elements_flow.index((elem_type, elem_idx, element))

        for j in range(current_idx_in_flow - 1, -1, -1):
            prev_type, prev_idx, prev_element = elements_flow[j]
            if prev_type != 'paragraph':
                break

            para_text = ''
            for t in prev_element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text:
                    para_text += t.text
            para_text = para_text.strip()

            if not para_text:
                continue

            match = re.match(r'^Таблица\s*(\d+)', para_text, re.IGNORECASE)
            if match:
                table_numbers.append(int(match.group(1)))
                break

            break

    figure_caption_pattern = re.compile(r'^Рис\.?\s*(\d+)', re.IGNORECASE)

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        match = figure_caption_pattern.match(text)
        if match:
            figure_numbers.append(int(match.group(1)))

    if table_numbers:
        expected_table_nums = list(range(1, len(table_numbers) + 1))
        for i, (actual, expected) in enumerate(zip(table_numbers, expected_table_nums)):
            if actual != expected:
                errors.append(ReportError(
                    id=f"Т-6-table-num-{i}",
                    code="Т-6",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=0,
                        structural_path=f"Таблица {i + 1}"
                    ),
                    fragment=f"Таблица {actual}",
                    rule="Нумерация таблиц должна быть сквозной: 1, 2, 3, ...",
                    rule_citation="§4.4, с. 50-52",
                    found_value=str(actual),
                    expected_value=str(expected),
                    recommendation=f"Исправьте номер таблицы на {expected}"
                ))
                break

    if figure_numbers:
        expected_figure_nums = list(range(1, len(figure_numbers) + 1))
        for i, (actual, expected) in enumerate(zip(figure_numbers, expected_figure_nums)):
            if actual != expected:
                errors.append(ReportError(
                    id=f"Т-6-figure-num-{i}",
                    code="Т-6",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=0,
                        structural_path=f"Рисунок {i + 1}"
                    ),
                    fragment=f"Рис. {actual}",
                    rule="Нумерация рисунков должна быть сквозной: 1, 2, 3, ...",
                    rule_citation="§4.4, с. 50-52",
                    found_value=str(actual),
                    expected_value=str(expected),
                    recommendation=f"Исправьте номер рисунка на {expected}"
                ))
                break

    # Т-12: десятичная запятая
    decimal_point_pattern = re.compile(r'\b\d+\.\d+\b')

    for table_idx, table in enumerate(tables_list):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue

                    matches = decimal_point_pattern.findall(text)
                    if matches:
                        for match in matches:
                            errors.append(ReportError(
                                id=f"Т-12-decimal-{table_idx}-{row_idx}-{cell_idx}",
                                code="Т-12",
                                type="formatting",
                                severity="error",
                                location=ErrorLocation(
                                    paragraph_index=0,
                                    structural_path=f"Таблица {table_idx + 1}, ячейка [{row_idx + 1}, {cell_idx + 1}]"
                                ),
                                fragment=text[:100],
                                rule="Дробные числа должны использовать запятую как десятичный разделитель",
                                rule_citation="§4.5, с. 51",
                                found_value=match,
                                expected_value=match.replace('.', ','),
                                recommendation="Замените точку на запятую в дробных числах"
                            ))

    return errors


def validate_references_format(doc: Document, rules: dict[str, Any]) -> list[ReportError]:
    """Проверяет список литературы (Л-1..Л-12)."""
    errors: list[ReportError] = []

    ref_section_paragraphs = []
    ref_section_start_idx = 0
    in_refs = False

    for para_idx, para in enumerate(doc.paragraphs):
        if "список литературы" in para.text.lower():
            in_refs = True
            ref_section_start_idx = para_idx + 1
            continue
        if in_refs and para.style and "Heading" in para.style.name:
            break
        if in_refs and para.text.strip():
            ref_section_paragraphs.append(para.text.strip())

    # Л-7
    min_sources = rules.get("references", {}).get("min_sources", 40)
    if len(ref_section_paragraphs) < min_sources:
        errors.append(ReportError(
            id="Л-7-count",
            code="Л-7",
            type="formatting",
            severity="error",
            location=ErrorLocation(
                paragraph_index=0,
                structural_path="Список литературы"
            ),
            fragment="Список литературы",
            rule=f"Список литературы должен содержать не менее {min_sources} источников",
            rule_citation="§3.7, с. 44",
            found_value=str(len(ref_section_paragraphs)),
            expected_value=str(min_sources),
            recommendation="Добавьте недостающие источники в список литературы"
        ))

    # Л-5
    numbering_pattern = re.compile(r'^(\d+)\.')
    source_numbers = []
    has_numbering = False

    for idx, para_text in enumerate(ref_section_paragraphs):
        match = numbering_pattern.match(para_text)
        if match:
            source_numbers.append(int(match.group(1)))
            has_numbering = True

    if not has_numbering and len(ref_section_paragraphs) > 0:
        errors.append(ReportError(
            id="Л-5-no-numbering",
            code="Л-5",
            type="formatting",
            severity="error",
            location=ErrorLocation(
                paragraph_index=ref_section_start_idx,
                structural_path="Список литературы"
            ),
            fragment=ref_section_paragraphs[0][:100] if ref_section_paragraphs else "Список литературы",
            rule="Источники в списке литературы должны иметь сквозную нумерацию: 1, 2, 3, ...",
            rule_citation="§4.5, с. 52",
            found_value="нумерация отсутствует",
            expected_value="1, 2, 3, ...",
            recommendation="Добавьте нумерацию к каждому источнику"
        ))
    elif source_numbers:
        expected_numbers = list(range(1, len(source_numbers) + 1))
        for i, (actual, expected) in enumerate(zip(source_numbers, expected_numbers)):
            if actual != expected:
                errors.append(ReportError(
                    id=f"Л-5-num-{i}",
                    code="Л-5",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=ref_section_start_idx + i,
                        structural_path="Список литературы"
                    ),
                    fragment=ref_section_paragraphs[i][:100],
                    rule="Нумерация источников должна быть сплошной: 1, 2, 3, ...",
                    rule_citation="§4.5, с. 52",
                    found_value=str(actual),
                    expected_value=str(expected),
                    recommendation=f"Исправьте номер источника на {expected}"
                ))
                break

    # Л-4
    def extract_surname(text: str) -> tuple[str, str]:
        clean_text = re.sub(r'^\d+\.\s*', '', text.strip())
        author_part = re.split(r'[,.]', clean_text)[0].strip()
        is_cyrillic = bool(re.search(r'[А-ЯЁа-яё]', author_part))
        return (author_part.lower(), 'cyrillic' if is_cyrillic else 'latin')

    def surname_sort_key(item: tuple[int, str]):
        idx, text = item
        surname, lang = extract_surname(text)
        lang_order = 0 if lang == 'cyrillic' else 1
        return (lang_order, surname)

    if len(ref_section_paragraphs) > 1:
        indexed_sources = list(enumerate(ref_section_paragraphs))
        sorted_sources = sorted(indexed_sources, key=surname_sort_key)

        for i, (orig_idx, _) in enumerate(indexed_sources):
            sorted_idx, _ = sorted_sources[i]
            if orig_idx != sorted_idx:
                curr_surname, curr_lang = extract_surname(ref_section_paragraphs[orig_idx])
                next_surname, next_lang = extract_surname(ref_section_paragraphs[sorted_idx])

                errors.append(ReportError(
                    id=f"Л-4-order-{i}",
                    code="Л-4",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=ref_section_start_idx + orig_idx,
                        structural_path="Список литературы"
                    ),
                    fragment=ref_section_paragraphs[orig_idx][:100],
                    rule="Источники должны быть расположены по алфавиту: сначала русские, затем иностранные",
                    rule_citation="§4.5, с. 52",
                    found_value=f"«{curr_surname}» после «{next_surname}»",
                    expected_value=f"«{next_surname}» перед «{curr_surname}»",
                    recommendation="Расположите источники в алфавитном порядке"
                ))
                break

    # Л-8
    current_year = datetime.now().year
    year_threshold = current_year - 10
    year_pattern = re.compile(r'\b(19|20)\d{2}\b')

    recent_sources_count = 0
    total_sources_with_year = 0

    for para_text in ref_section_paragraphs:
        year_match = year_pattern.search(para_text)
        if year_match:
            total_sources_with_year += 1
            year = int(year_match.group(0))
            if year >= year_threshold:
                recent_sources_count += 1

    if total_sources_with_year > 0:
        recent_percentage = (recent_sources_count / total_sources_with_year) * 100
        if recent_percentage < 70:
            errors.append(ReportError(
                id="Л-8-recency",
                code="Л-8",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=ref_section_start_idx,
                    structural_path="Список литературы"
                ),
                fragment="Список литературы",
                rule=f"Не менее 70% источников должны быть опубликованы в последние 10 лет",
                rule_citation="§4.5, с. 52",
                found_value=f"{recent_percentage:.1f}%",
                expected_value=">= 70%",
                recommendation="Добавьте более свежие источники"
            ))

    # ---------------------------------------------------------------------------
    # ИСПРАВЛЕНИЕ №2 (Л-9): детектируем записи с инициалами перед фамилией
    # ---------------------------------------------------------------------------
    # Паттерн для записей которые НАЧИНАЮТСЯ с фамилии «Фамилия, И. О.»
    correct_author_pattern = re.compile(
        r'[А-ЯЁ][а-яё]+\s*,\s*[А-ЯЁ]\.\s*[А-ЯЁ]\.'     # кириллица
        r'|[A-Z][a-z]+\s*,\s*[A-Z]\.\s*[A-Z]\.'           # латиница
    )
    # Паттерн для записей которые начинаются с инициалов «И. И. Фамилия» — нарушение
    wrong_initials_first_pattern = re.compile(
        r'^(\d+\.\s*)?[А-ЯЁ]\.\s*[А-ЯЁ]\.\s*[А-ЯЁ][а-яё]'   # кириллица: И. И. Фамил...
        r'|^(\d+\.\s*)?[A-Z]\.\s*[A-Z]\.\s*[A-Z][a-z]'         # латиница
    )

    for idx, para_text in enumerate(ref_section_paragraphs):
        # Случай 1: запись начинается с инициалов вместо фамилии
        if wrong_initials_first_pattern.match(para_text):
            errors.append(ReportError(
                id=f"Л-9-initials-first-{idx}",
                code="Л-9",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=ref_section_start_idx + idx,
                    structural_path="Список литературы"
                ),
                fragment=para_text[:100],
                rule="Автор должен быть указан в формате: Фамилия, И. О.",
                rule_citation="§4.5, с. 52",
                found_value=para_text[:50],
                expected_value="Фамилия, И. О.",
                recommendation="Исправьте: сначала фамилия, затем инициалы через запятую"
            ))
            continue

        # Случай 2: запись начинается с фамилии, но без запятой перед инициалами
        # Паттерн «Фамилия И. О.» (без запятой) — нарушение
        no_comma_pattern = re.compile(
            r'^(\d+\.\s*)?([А-ЯЁ][а-яё]{2,})\s+[А-ЯЁ]\.'
        )
        if no_comma_pattern.match(para_text) and not correct_author_pattern.search(para_text):
            errors.append(ReportError(
                id=f"Л-9-no-comma-{idx}",
                code="Л-9",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=ref_section_start_idx + idx,
                    structural_path="Список литературы"
                ),
                fragment=para_text[:100],
                rule="Автор должен быть указан в формате: Фамилия, И. О.",
                rule_citation="§4.5, с. 52",
                found_value=para_text[:50],
                expected_value="Фамилия, И. О.",
                recommendation="Добавьте запятую после фамилии автора"
            ))

    # Л-10
    url_pattern = re.compile(r'https?://[^\s]+')
    access_date_pattern = re.compile(r'\(дата обращения:\s*\d{2}\.\d{2}\.\d{4}\)')

    for idx, para_text in enumerate(ref_section_paragraphs):
        if url_pattern.search(para_text):
            if not access_date_pattern.search(para_text):
                errors.append(ReportError(
                    id=f"Л-10-url-{idx}",
                    code="Л-10",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=ref_section_start_idx + idx,
                        structural_path="Список литературы"
                    ),
                    fragment=para_text[:100],
                    rule="Для URL-источников должна быть указана дата обращения",
                    rule_citation="§4.5, с. 52",
                    found_value="URL без даты обращения",
                    expected_value="(дата обращения: ДД.ММ.ГГГГ)",
                    recommendation="Добавьте дату обращения после URL"
                ))

    # Л-11
    valid_source_numbers = set(source_numbers)
    ref_pattern = re.compile(r'\[(\d+)(?:\s*,\s*с\.\s*\d+)?\]')

    for para_idx, para in enumerate(doc.paragraphs):
        for match in ref_pattern.finditer(para.text):
            ref_number = int(match.group(1))
            if ref_number not in valid_source_numbers:
                errors.append(ReportError(
                    id=f"Л-11-ref-{para_idx}-{ref_number}",
                    code="Л-11",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path=f"Абзац {para_idx + 1}"
                    ),
                    fragment=match.group(0),
                    rule="Ссылка в тексте должна соответствовать источнику в списке литературы",
                    rule_citation="§4.3, с. 49",
                    found_value=f"№{ref_number}",
                    expected_value=f"номер от 1 до {len(valid_source_numbers)}",
                    recommendation=f"Исправьте ссылку или добавьте источник №{ref_number}"
                ))

    # Л-12
    hyphen_as_dash_pattern = re.compile(r'\s-\s|\d-\d')

    for idx, para_text in enumerate(ref_section_paragraphs):
        if hyphen_as_dash_pattern.search(para_text):
            errors.append(ReportError(
                id=f"Л-12-hyphen-{idx}",
                code="Л-12",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=ref_section_start_idx + idx,
                    structural_path="Список литературы"
                ),
                fragment=para_text[:100],
                rule="В библиографии должны использоваться длинные тире «–» (U+2013)",
                rule_citation="§4.5, с. 52",
                found_value="-",
                expected_value="–",
                recommendation="Замените дефис на длинное тире"
            ))

    # Л-1
    inline_pattern = re.compile(r'\(\d+\)')
    for para_idx, para in enumerate(doc.paragraphs):
        matches = inline_pattern.findall(para.text)
        if matches:
            errors.append(ReportError(
                id=f"Л-1-{para_idx}",
                code="Л-1",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=para_idx,
                    structural_path=f"Абзац {para_idx + 1}"
                ),
                fragment=para.text[:100],
                rule="Ссылки оформляются в квадратных скобках: [N] или [N, с. X]",
                rule_citation="§4.3, с. 49",
                found_value=matches[0],
                expected_value="[N] или [N, с. X]",
                recommendation="Замените круглые скобки на квадратные"
            ))

    # Л-3
    multi_pattern = re.compile(r'\[(\d+(?:;\s*\d+)+)\]')
    for para_idx, para in enumerate(doc.paragraphs):
        for match in multi_pattern.finditer(para.text):
            nums = [int(n.strip()) for n in match.group(1).split(';')]
            if nums != sorted(nums):
                errors.append(ReportError(
                    id=f"Л-3-{para_idx}",
                    code="Л-3",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path=f"Абзац {para_idx + 1}"
                    ),
                    fragment=match.group(0),
                    rule="Несколько источников указываются в арифметическом порядке через ';'",
                    rule_citation="§4.3, с. 49",
                    found_value=match.group(0),
                    expected_value=f"[{'; '.join(str(n) for n in sorted(nums))}]",
                    recommendation="Расположите номера источников в порядке возрастания"
                ))

    return errors


def validate_volume(doc: Document, rules: dict[str, Any]) -> list[ReportError]:
    errors: list[ReportError] = []

    volume_rules = rules.get("volume", {})
    total_min = volume_rules.get("total_chars_min", 90000)
    total_max = volume_rules.get("total_chars_max", 108000)
    theory_min = volume_rules.get("theory_chapter_chars_min", 27000)
    theory_max = volume_rules.get("theory_chapter_chars_max", 36000)
    empirical_min = volume_rules.get("empirical_chapter_chars_min", 45000)
    empirical_max = volume_rules.get("empirical_chapter_chars_max", 54000)

    heading_styles = {"Heading 1", "Heading 2"}
    current_section = None
    section_texts: dict[str, list[str]] = {}
    all_text_parts: list[str] = []

    for para in doc.paragraphs:
        if para.style and para.style.name in heading_styles:
            title = para.text.strip()
            if title.lower().startswith("глава 1"):
                current_section = "Глава 1"
            elif title.lower().startswith("глава 2"):
                current_section = "Глава 2"
            elif title.lower() in ["введение", "заключение", "список литературы", "содержание"]:
                current_section = title.lower()
            else:
                current_section = title
        elif para.text.strip():
            all_text_parts.append(para.text)
            if current_section:
                if current_section not in section_texts:
                    section_texts[current_section] = []
                section_texts[current_section].append(para.text)

    total_text = "".join(all_text_parts)
    total_chars = len(total_text)

    if total_chars < total_min:
        errors.append(ReportError(
            id="Ф-11-below-min",
            code="Ф-11",
            type="formatting",
            severity="error",
            location=ErrorLocation(
                paragraph_index=0,
                structural_path="Документ целиком"
            ),
            fragment=f"Общий объём: {total_chars} знаков",
            rule=f"Объём ВКР должен быть от {total_min} до {total_max} знаков с пробелами",
            rule_citation="§4.1, с. 46",
            found_value=str(total_chars),
            expected_value=f"{total_min}-{total_max}",
            recommendation=f"Добавьте текст. Не хватает {total_min - total_chars} знаков."
        ))
    elif total_chars > total_max:
        errors.append(ReportError(
            id="Ф-11-above-max",
            code="Ф-11",
            type="formatting",
            severity="error",
            location=ErrorLocation(
                paragraph_index=0,
                structural_path="Документ целиком"
            ),
            fragment=f"Общий объём: {total_chars} знаков",
            rule=f"Объём ВКР должен быть от {total_min} до {total_max} знаков с пробелами",
            rule_citation="§4.1, с. 46",
            found_value=str(total_chars),
            expected_value=f"{total_min}-{total_max}",
            recommendation=f"Сократите текст. Превышение на {total_chars - total_max} знаков."
        ))

    chapter1_text = "".join(section_texts.get("Глава 1", []))
    chapter1_chars = len(chapter1_text)

    if chapter1_chars > 0:
        if chapter1_chars < theory_min:
            errors.append(ReportError(
                id="Ф-12-chapter1-below-min",
                code="Ф-12",
                type="formatting",
                severity="error",
                location=ErrorLocation(paragraph_index=0, structural_path="Глава 1"),
                fragment=f"Глава 1: {chapter1_chars} знаков",
                rule=f"Теоретическая глава должна содержать от {theory_min} до {theory_max} знаков",
                rule_citation="§3.4, с. 23",
                found_value=str(chapter1_chars),
                expected_value=f"{theory_min}-{theory_max}",
                recommendation=f"Расширьте Главу 1. Не хватает {theory_min - chapter1_chars} знаков."
            ))
        elif chapter1_chars > theory_max:
            errors.append(ReportError(
                id="Ф-12-chapter1-above-max",
                code="Ф-12",
                type="formatting",
                severity="error",
                location=ErrorLocation(paragraph_index=0, structural_path="Глава 1"),
                fragment=f"Глава 1: {chapter1_chars} знаков",
                rule=f"Теоретическая глава должна содержать от {theory_min} до {theory_max} знаков",
                rule_citation="§3.4, с. 23",
                found_value=str(chapter1_chars),
                expected_value=f"{theory_min}-{theory_max}",
                recommendation=f"Сократите Главу 1. Превышение на {chapter1_chars - theory_max} знаков."
            ))

    chapter2_text = "".join(section_texts.get("Глава 2", []))
    chapter2_chars = len(chapter2_text)

    if chapter2_chars > 0:
        if chapter2_chars < empirical_min:
            errors.append(ReportError(
                id="Ф-13-chapter2-below-min",
                code="Ф-13",
                type="formatting",
                severity="error",
                location=ErrorLocation(paragraph_index=0, structural_path="Глава 2"),
                fragment=f"Глава 2: {chapter2_chars} знаков",
                rule=f"Эмпирическая глава должна содержать от {empirical_min} до {empirical_max} знаков",
                rule_citation="§3.5, с. 30",
                found_value=str(chapter2_chars),
                expected_value=f"{empirical_min}-{empirical_max}",
                recommendation=f"Расширьте Главу 2. Не хватает {empirical_min - chapter2_chars} знаков."
            ))
        elif chapter2_chars > empirical_max:
            errors.append(ReportError(
                id="Ф-13-chapter2-above-max",
                code="Ф-13",
                type="formatting",
                severity="error",
                location=ErrorLocation(paragraph_index=0, structural_path="Глава 2"),
                fragment=f"Глава 2: {chapter2_chars} знаков",
                rule=f"Эмпирическая глава должна содержать от {empirical_min} до {empirical_max} знаков",
                rule_citation="§3.5, с. 30",
                found_value=str(chapter2_chars),
                expected_value=f"{empirical_min}-{empirical_max}",
                recommendation=f"Сократите Главу 2. Превышение на {chapter2_chars - empirical_max} знаков."
            ))

    return errors


def validate_typography_format(doc: Document, rules: dict[str, Any]) -> list[ReportError]:
    """
    Проверяет типографику текста.

    Н-2: пробелы между инициалами
    Н-4: кавычки-лапки
    Н-5: тире вместо дефиса между датами/числами
    Н-6: аббревиатуры без расшифровки
    Н-7: автоматическая нумерация списков
    """
    errors: list[ReportError] = []

    no_space_pattern = re.compile(r'[А-ЯЁ]\.[А-ЯЁ]\.[А-ЯЁ][а-яё]+')

    # ---------------------------------------------------------------------------
    # ИСПРАВЛЕНИЕ №3 (Н-4): ловим все виды «неправильных» кавычек:
    #   ASCII: "..."  (U+0022 пары)
    #   Немецкие: „..." (U+201E + U+201C) — Word может вставить автозаменой
    #   Одиночная прямая: '...'
    # ---------------------------------------------------------------------------
    wrong_quotes = re.compile(
        r'"[^"]*"'                      # ASCII "..."
        r'|\u201e[^\u201c]*\u201c'      # „..." немецкие (открыв. снизу)
        r'|\u201c[^\u201d]*\u201d'      # "..." левая+правая (типографские англ.)
    )

    abbrev_pattern = re.compile(r'(?<![а-яёa-z])[А-ЯЁ]{2,}(?![а-яё])')
    explained_pattern = re.compile(r'\([А-ЯЁ]{2,}\)')

    # ---------------------------------------------------------------------------
    # ИСПРАВЛЕНИЕ №4 (Н-5): более точный паттерн — дефис между цифрами
    # без длинного тире (–) рядом
    # ---------------------------------------------------------------------------
    hyphen_between_numbers = re.compile(r'(\d)-(\d)')   # убрали \s* — строже

    manual_numbering_pattern = re.compile(r'^\s*\d+\.\s+')
    bullet_markers = ['•', '-', '◦', '▪', '‣', '⁃']

    found_abbrevs: set[str] = set()
    in_list = False
    list_marker_type = None

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text

        # Н-2
        match_no_space = no_space_pattern.search(text)
        if match_no_space:
            errors.append(ReportError(
                id=f"Н-2-{para_idx}",
                code="Н-2",
                type="style",
                severity="warning",
                location=ErrorLocation(
                    paragraph_index=para_idx,
                    structural_path=f"Абзац {para_idx + 1}"
                ),
                fragment=text[:100],
                rule="Между инициалами и перед фамилией должен быть пробел: И. И. Иванов",
                rule_citation="§4.2, с. 48",
                found_value=match_no_space.group(0),
                expected_value="И. И. Иванов",
                recommendation="Добавьте пробелы между инициалами"
            ))

        # Н-4
        if wrong_quotes.search(text):
            errors.append(ReportError(
                id=f"Н-4-{para_idx}",
                code="Н-4",
                type="style",
                severity="warning",
                location=ErrorLocation(
                    paragraph_index=para_idx,
                    structural_path=f"Абзац {para_idx + 1}"
                ),
                fragment=text[:100],
                rule="Кавычки должны быть угловыми: «текст»",
                rule_citation="§4.2, с. 48",
                found_value='"..."',
                expected_value='«...»',
                recommendation='Замените кавычки "..." на «...»'
            ))

        # Н-5
        hyphen_match = hyphen_between_numbers.search(text)
        if hyphen_match:
            errors.append(ReportError(
                id=f"Н-5-{para_idx}",
                code="Н-5",
                type="style",
                severity="warning",
                location=ErrorLocation(
                    paragraph_index=para_idx,
                    structural_path=f"Абзац {para_idx + 1}"
                ),
                fragment=text[:100],
                rule="Между числами/датами должно использоваться тире (–), а не дефис (-)",
                rule_citation="§4.2, с. 48",
                found_value=hyphen_match.group(0),
                expected_value="число – число",
                recommendation="Замените дефис на тире в диапазоне чисел/дат"
            ))

        # ---------------------------------------------------------------------------
        # ИСПРАВЛЕНИЕ №5 (Н-7): добавлена защита от pPr = None
        # ---------------------------------------------------------------------------
        if manual_numbering_pattern.match(text):
            pPr = para._p.pPr   # может быть None
            is_numbered = False
            if pPr is not None:           # ← была ошибка: pPr.find без проверки
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    is_numbered = True

            if not is_numbered:
                errors.append(ReportError(
                    id=f"Н-7-manual-{para_idx}",
                    code="Н-7",
                    type="style",
                    severity="warning",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path=f"Абзац {para_idx + 1}"
                    ),
                    fragment=text[:100],
                    rule="Списки должны использовать автоматическую нумерацию, а не ручную",
                    rule_citation="§4.2, с. 48",
                    found_value="ручная нумерация",
                    expected_value="автоматическая нумерация",
                    recommendation="Используйте автоматическую нумерацию для списков"
                ))

        # Н-7: смешанные маркеры
        stripped_text = text.strip()
        current_marker = None
        for marker in bullet_markers:
            if stripped_text.startswith(marker):
                current_marker = marker
                break

        if current_marker is not None:
            if not in_list:
                in_list = True
                list_marker_type = current_marker
            elif list_marker_type != current_marker:
                errors.append(ReportError(
                    id=f"Н-7-mixed-{para_idx}",
                    code="Н-7",
                    type="style",
                    severity="warning",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path=f"Абзац {para_idx + 1}"
                    ),
                    fragment=text[:100],
                    rule="В одном списке должны использоваться унифицированные маркеры",
                    rule_citation="§4.2, с. 48",
                    found_value=f"маркер '{current_marker}'",
                    expected_value=f"единый маркер '{list_marker_type}'",
                    recommendation="Используйте одинаковые маркеры во всём списке"
                ))
        else:
            in_list = False
            list_marker_type = None

        # ---------------------------------------------------------------------------
        # ИСПРАВЛЕНИЕ №6 (Н-6): расшифровки собираются ДО проверки аббревиатур
        # (порядок уже был верным, но теперь явно отделён блок)
        # ---------------------------------------------------------------------------
        for m in explained_pattern.finditer(text):
            found_abbrevs.add(m.group(0)[1:-1])

        for m in abbrev_pattern.finditer(text):
            abbrev = m.group(0)
            if abbrev not in found_abbrevs:
                errors.append(ReportError(
                    id=f"Н-6-{para_idx}-{abbrev}",
                    code="Н-6",
                    type="style",
                    severity="warning",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path=f"Абзац {para_idx + 1}"
                    ),
                    fragment=text[:100],
                    rule="При первом использовании аббревиатуры дайте расшифровку: полное название (АБВ)",
                    rule_citation="§4.1, с. 46",
                    found_value=abbrev,
                    expected_value=f"полное название ({abbrev})",
                    recommendation=f"Расшифруйте аббревиатуру {abbrev} при первом использовании"
                ))
                found_abbrevs.add(abbrev)

    return errors


def validate_toc(doc: Document, rules: dict[str, Any]) -> list[ReportError]:
    errors: list[ReportError] = []

    SKIP_IN_TOC_CHECK = {"содержани", "оглавлени"}
    headings_to_check: list[tuple[int, str]] = []
    for para_idx, para in enumerate(doc.paragraphs):
        if para.style and para.style.name in ("Heading 1", "Heading 2"):
            title = para.text.strip()
            if not title:
                continue
            title_lower = title.lower()
            if any(s in title_lower for s in SKIP_IN_TOC_CHECK):
                continue
            headings_to_check.append((para_idx, title))

    toc_start_idx = None
    toc_end_idx = None

    for para_idx, para in enumerate(doc.paragraphs):
        text_lower = para.text.lower()
        if any(s in text_lower for s in SKIP_IN_TOC_CHECK):
            toc_start_idx = para_idx
            continue
        if toc_start_idx is not None and toc_end_idx is None:
            if para.style and para.style.name == "Heading 1":
                toc_end_idx = para_idx
                break

    if toc_start_idx is None:
        if headings_to_check:
            errors.append(ReportError(
                id="Со-1-no-toc",
                code="Со-1",
                type="formatting",
                severity="error",
                location=ErrorLocation(paragraph_index=0, structural_path="Структура документа"),
                fragment="Содержание отсутствует",
                rule="Документ должен содержать раздел «Содержание» со всеми заголовками",
                rule_citation="§3.2, с. 12",
                found_value="раздел Содержание отсутствует",
                expected_value="раздел Содержание с перечнем всех заголовков",
                recommendation="Добавьте раздел «Содержание» после титульного листа"
            ))
        return errors

    end = toc_end_idx if toc_end_idx else len(doc.paragraphs)
    toc_paragraphs = [
        doc.paragraphs[i].text.strip()
        for i in range(toc_start_idx + 1, end)
        if doc.paragraphs[i].text.strip()
    ]

    has_toc_field = False
    for para in doc.paragraphs[toc_start_idx:end]:
        xml = para._p.xml
        if 'TOC' in xml or 'w:fldChar' in xml:
            has_toc_field = True
            break

    if has_toc_field:
        return errors

    if not toc_paragraphs:
        errors.append(ReportError(
            id="Со-1-empty-toc",
            code="Со-1",
            type="formatting",
            severity="error",
            location=ErrorLocation(paragraph_index=toc_start_idx, structural_path="Содержание"),
            fragment="Содержание пустое",
            rule="Содержание должно отражать все заголовки с номерами страниц",
            rule_citation="§3.2, с. 12",
            found_value="содержание пустое",
            expected_value="перечень всех заголовков",
            recommendation="Заполните содержание или используйте автоматическое оглавление Word"
        ))
        return errors

    def heading_in_toc(title: str, toc_lines: list[str]) -> bool:
        title_norm = re.sub(r'\s+', ' ', title.lower().strip())
        for line in toc_lines:
            if title_norm in re.sub(r'\s+', ' ', line.lower()):
                return True
        words = [w for w in title_norm.split() if len(w) > 3]
        if not words:
            return True
        for line in toc_lines:
            line_norm = re.sub(r'\s+', ' ', line.lower())
            matches = sum(1 for w in words if w in line_norm)
            if matches / len(words) >= 0.7:
                return True
        return False

    for heading_idx, title in headings_to_check:
        if not heading_in_toc(title, toc_paragraphs):
            errors.append(ReportError(
                id=f"Со-1-{heading_idx}",
                code="Со-1",
                type="formatting",
                severity="error",
                location=ErrorLocation(
                    paragraph_index=heading_idx,
                    structural_path=f"Заголовок «{title[:50]}»"
                ),
                fragment=title[:100],
                rule="Все заголовки должны быть отражены в содержании",
                rule_citation="§3.2, с. 12",
                found_value=f"«{title}» не найден в содержании",
                expected_value="заголовок присутствует в содержании",
                recommendation="Добавьте этот заголовок в содержание"
            ))

    return errors


def validate_appendix(doc: Document, rules: dict[str, Any]) -> list[ReportError]:
    """П-1..П-4: оформление приложений."""
    errors: list[ReportError] = []

    appendix_heading_pattern = re.compile(
        r'^Приложение\s+([А-ЯЁA-Z\d])\s*$', re.IGNORECASE
    )
    appendix_ref_pattern = re.compile(
        r'(?:прил\.\s*([А-ЯЁA-Z\d]+)|\bприложени[еяю]\s+([А-ЯЁA-Z\d]+))',
        re.IGNORECASE
    )

    def _has_page_break_before(para, para_idx: int) -> bool:
        pPr = para._p.find(qn('w:pPr'))
        if pPr is not None:
            pb = pPr.find(qn('w:pageBreakBefore'))
            if pb is not None:
                val = pb.get(qn('w:val'))
                if val is None or val in ('1', 'true', 'on'):
                    return True
        if para_idx > 0:
            prev_para = doc.paragraphs[para_idx - 1]
            for br in prev_para._p.iter(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    return True
        return False

    def _get_alignment(para) -> str | None:
        pPr = para._p.find(qn('w:pPr'))
        jc_el = pPr.find(qn('w:jc')) if pPr is not None else None
        if jc_el is not None:
            return jc_el.get(qn('w:val'))
        if para.style and para.style.paragraph_format:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            fmt = para.style.paragraph_format
            mapping = {
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.LEFT: "left",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
            }
            if fmt.alignment in mapping:
                return mapping[fmt.alignment]
        return None

    refs_order: list[str] = []
    seen_refs: set[str] = set()
    for para in doc.paragraphs:
        if appendix_heading_pattern.match(para.text.strip()):
            continue
        for m in appendix_ref_pattern.finditer(para.text):
            letter = (m.group(1) or m.group(2) or "").upper().strip()
            if letter and letter not in seen_refs:
                refs_order.append(letter)
                seen_refs.add(letter)

    appendices: list[dict] = []

    i = 0
    paras = doc.paragraphs
    while i < len(paras):
        para = paras[i]
        m = appendix_heading_pattern.match(para.text.strip())
        if m:
            letter = m.group(1).upper()
            app = {"idx": i, "letter": letter, "title_idx": None, "title": None}

            if not _has_page_break_before(para, i):
                errors.append(ReportError(
                    id=f"П-1-{i}",
                    code="П-1",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(paragraph_index=i,
                                           structural_path=f"Приложение {letter}"),
                    fragment=para.text.strip()[:100],
                    rule="Каждое приложение должно начинаться с новой страницы",
                    rule_citation="§4.6, с. 59",
                    found_value="нет разрыва страницы",
                    expected_value="разрыв страницы перед приложением",
                    recommendation="Поставьте курсор перед «Приложение» и нажмите Ctrl+Enter"
                ))

            alignment = _get_alignment(para)
            if alignment != "right":
                errors.append(ReportError(
                    id=f"П-2-{i}",
                    code="П-2",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(paragraph_index=i,
                                           structural_path=f"Приложение {letter}"),
                    fragment=para.text.strip()[:100],
                    rule="Надпись «Приложение N» должна быть выровнена по правому краю",
                    rule_citation="§4.6, с. 59",
                    found_value=alignment or "не задано",
                    expected_value="right",
                    recommendation="Выделите строку «Приложение N» → Главная → По правому краю"
                ))

            j = i + 1
            while j < len(paras) and not paras[j].text.strip():
                j += 1

            if j < len(paras):
                next_para = paras[j]
                next_text = next_para.text.strip()
                is_another_app = bool(appendix_heading_pattern.match(next_text))
                is_h1 = next_para.style and next_para.style.name == "Heading 1"

                if not is_another_app and not is_h1 and next_text:
                    app["title_idx"] = j
                    app["title"] = next_text

                    title_align = _get_alignment(next_para)
                    if title_align != "center":
                        errors.append(ReportError(
                            id=f"П-3-align-{j}",
                            code="П-3",
                            type="formatting",
                            severity="error",
                            location=ErrorLocation(paragraph_index=j,
                                                   structural_path=f"Название приложения {letter}"),
                            fragment=next_text[:100],
                            rule="Название приложения должно быть выровнено по центру",
                            rule_citation="§4.6, с. 59",
                            found_value=title_align or "не задано",
                            expected_value="center",
                            recommendation="Выделите название → Главная → По центру"
                        ))

                    if next_text.endswith('.'):
                        errors.append(ReportError(
                            id=f"П-3-dot-{j}",
                            code="П-3",
                            type="formatting",
                            severity="error",
                            location=ErrorLocation(paragraph_index=j,
                                                   structural_path=f"Название приложения {letter}"),
                            fragment=next_text[:100],
                            rule="Название приложения не должно заканчиваться точкой",
                            rule_citation="§4.6, с. 59",
                            found_value=next_text[-10:],
                            expected_value="без точки в конце",
                            recommendation="Удалите точку в конце названия приложения"
                        ))

            appendices.append(app)
        i += 1

    if refs_order and appendices:
        app_letters_in_doc = [a["letter"] for a in appendices]
        filtered_refs = [r for r in refs_order if r in app_letters_in_doc]

        if filtered_refs and filtered_refs != app_letters_in_doc[:len(filtered_refs)]:
            errors.append(ReportError(
                id="П-4-order",
                code="П-4",
                type="formatting",
                severity="error",
                location=ErrorLocation(paragraph_index=appendices[0]["idx"],
                                       structural_path="Приложения"),
                fragment=f"Порядок в документе: {', '.join(app_letters_in_doc)}",
                rule="Приложения нумеруются в порядке упоминания в тексте",
                rule_citation="§4.6, с. 59",
                found_value=f"в документе: {', '.join(app_letters_in_doc)}",
                expected_value=f"по порядку ссылок: {', '.join(filtered_refs)}",
                recommendation="Переставьте приложения в порядке их упоминания в тексте"
            ))

        missing_apps = [r for r in refs_order if r not in app_letters_in_doc]
        if missing_apps:
            errors.append(ReportError(
                id="П-4-missing",
                code="П-4",
                type="formatting",
                severity="error",
                location=ErrorLocation(paragraph_index=appendices[0]["idx"],
                                       structural_path="Приложения"),
                fragment=f"Отсутствуют приложения: {', '.join(missing_apps)}",
                rule="Все упомянутые в тексте приложения должны присутствовать в документе",
                rule_citation="§4.6, с. 59",
                found_value=f"приложения {', '.join(missing_apps)} отсутствуют",
                expected_value=f"приложения {', '.join(missing_apps)} должны быть в документе",
                recommendation="Добавьте отсутствующие приложения или исправьте ссылки в тексте"
            ))

    return errors


def validate_repeated_references(doc: Document, rules: dict[str, Any]) -> list[ReportError]:
    """Л-2: повторная ссылка — [там же, с. X]."""
    errors: list[ReportError] = []

    ref_with_page = re.compile(r'\[(\d+),\s*с\.\s*(\d+)\]')
    correct_repeat = re.compile(r'\[там же(?:,\s*с\.\s*\d+)?\]', re.IGNORECASE)

    paras = doc.paragraphs

    for para_idx, para in enumerate(paras):
        text = para.text

        refs_in_para: dict[str, list[str]] = {}
        for m in ref_with_page.finditer(text):
            src = m.group(1)
            page = m.group(2)
            refs_in_para.setdefault(src, []).append(page)

        for src, pages in refs_in_para.items():
            if len(pages) > 1 and not correct_repeat.search(text):
                errors.append(ReportError(
                    id=f"Л-2-same-para-{para_idx}-{src}",
                    code="Л-2",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path=f"Абзац {para_idx + 1}"
                    ),
                    fragment=text[:100],
                    rule="Повторная ссылка на тот же источник в одном абзаце должна быть [там же, с. X]",
                    rule_citation="§4.3, с. 49",
                    found_value=f"[{src}, с. {pages[0]}]...[{src}, с. {pages[1]}]",
                    expected_value=f"[{src}, с. {pages[0]}]...[там же, с. {pages[1]}]",
                    recommendation="Замените второй [N, с. X] на [там же, с. X]"
                ))

        if para_idx + 1 >= len(paras):
            continue

        last_ref_in_current = None
        for m in ref_with_page.finditer(text):
            last_ref_in_current = m.group(1)

        if last_ref_in_current is None:
            continue

        next_para = paras[para_idx + 1]
        next_text = next_para.text

        first_ref_match = re.match(r'^\s*\[(\d+),\s*с\.\s*\d+\]', next_text)
        if first_ref_match:
            first_src_next = first_ref_match.group(1)
            if (first_src_next == last_ref_in_current
                    and not correct_repeat.search(next_text)):
                errors.append(ReportError(
                    id=f"Л-2-next-para-{para_idx + 1}-{first_src_next}",
                    code="Л-2",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=para_idx + 1,
                        structural_path=f"Абзац {para_idx + 2}"
                    ),
                    fragment=next_text[:100],
                    rule="Повторная ссылка в следующем абзаце должна быть [там же, с. X]",
                    rule_citation="§4.3, с. 49",
                    found_value=first_ref_match.group(0),
                    expected_value="[там же, с. X]",
                    recommendation="Замените ссылку на [там же, с. X] в начале абзаца"
                ))

    return errors


def validate_list_numbering(doc: Document, rules: dict[str, Any]) -> list[ReportError]:
    """Л-5: сплошная нумерация источников."""
    errors: list[ReportError] = []

    bibliography_start_idx = None
    bibliography_end_idx = None

    for para_idx, para in enumerate(doc.paragraphs):
        text_lower = para.text.lower()
        if "список литературы" in text_lower or "библиографический список" in text_lower:
            bibliography_start_idx = para_idx
            continue

        if bibliography_start_idx is not None and bibliography_end_idx is None:
            if para.style and para.style.name == "Heading 1":
                bibliography_end_idx = para_idx
                break

    if bibliography_start_idx is None:
        return errors

    end = bibliography_end_idx if bibliography_end_idx else len(doc.paragraphs)
    bibliography_paragraphs = []
    for i in range(bibliography_start_idx + 1, end):
        para = doc.paragraphs[i]
        if para.text.strip():
            bibliography_paragraphs.append((i, para))

    if not bibliography_paragraphs:
        return errors

    numbering_pattern = re.compile(r'^(\d+)[\.\)]\s')

    expected_number = 1

    for para_idx, para in bibliography_paragraphs:
        text = para.text.strip()

        match = numbering_pattern.match(text)
        if match:
            actual_number = int(match.group(1))
            if actual_number != expected_number:
                errors.append(ReportError(
                    id=f"Л-5-{para_idx}",
                    code="Л-5",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path="Список литературы"
                    ),
                    fragment=text[:100],
                    rule="Источники должны быть пронумерованы сплошной нумерацией",
                    rule_citation="§4.3, с. 49",
                    found_value=f"номер {actual_number}",
                    expected_value=f"номер {expected_number}",
                    recommendation="Исправьте нумерацию источников на сплошную"
                ))
            expected_number = actual_number + 1
        else:
            pPr = para._p.pPr
            has_indent = False
            if pPr is not None:
                ind_el = pPr.find(qn('w:ind'))
                if ind_el is not None:
                    left = ind_el.get(qn('w:left'))
                    if left is not None and int(left) > 0:
                        has_indent = True

            if not has_indent and len(text) > 10:
                errors.append(ReportError(
                    id=f"Л-5-no-num-{para_idx}",
                    code="Л-5",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(
                        paragraph_index=para_idx,
                        structural_path="Список литературы"
                    ),
                    fragment=text[:100],
                    rule="Источники должны быть пронумерованы сплошной нумерацией",
                    rule_citation="§4.3, с. 49",
                    found_value="нумерация отсутствует",
                    expected_value=f"номер {expected_number}.",
                    recommendation="Добавьте номер источника"
                ))

    return errors


def validate_format(docx_path: str, rules: dict[str, Any]) -> ValidationReport:
    """Полная валидация DOCX-документа."""
    doc = Document(docx_path)
    errors: list[ReportError] = []

    errors.extend(check_font_formatting(doc, rules))
    errors.extend(check_paragraph_formatting(doc, rules))
    errors.extend(check_margins(doc, rules))
    errors.extend(validate_structure(doc, rules))
    errors.extend(validate_tables(doc, rules))
    errors.extend(validate_references_format(doc, rules))
    errors.extend(validate_typography_format(doc, rules))
    errors.extend(validate_toc(doc, rules))
    errors.extend(validate_appendix(doc, rules))
    errors.extend(validate_repeated_references(doc, rules))
    errors.extend(validate_list_numbering(doc, rules))
    errors.extend(validate_volume(doc, rules))

    formatting_count = sum(1 for e in errors if e.type == "formatting")
    style_count = sum(1 for e in errors if e.type == "style")
    citation_count = sum(1 for e in errors if e.type == "citation_check")

    return ValidationReport(
        doc_id=str(uuid.uuid4()),
        created_at=datetime.now(timezone.utc),
        session_expires_at=datetime.now(timezone.utc) + timedelta(hours=1),
        summary=ReportSummary(
            total_errors=len(errors),
            formatting=formatting_count,
            style=style_count,
            citations=citation_count
        ),
        errors=errors
    )