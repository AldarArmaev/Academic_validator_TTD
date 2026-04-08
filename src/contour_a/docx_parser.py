"""
Парсер документов формата DOCX с извлечением форматирования.
Использует библиотеку python-docx для анализа структуры документа.
"""
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .document_loader import (
    Paragraph,
    Heading,
    Table,
    TableCell,
    Figure,
    Formula,
    Reference,
    ReferenceType,
    BibliographyEntry,
    Appendix,
    DocumentMetadata,
    DocumentContent,
)


class DOCXParser:
    """
    Парсер документов DOCX с извлечением текста, стилей и форматирования.
    
    Извлекает:
    - Абзацы с полным форматированием (шрифт, размер, выравнивание, отступы)
    - Заголовки с уровнями и нумерацией
    - Таблицы с подписями и ячейками
    - Рисунки и формулы
    - Библиографические ссылки
    - Нумерацию страниц
    """
    
    # Константы для конвертации единиц
    TWIPS_TO_PT = 20  # 1 пункт = 20 twips
    TWIPS_TO_LINE_FACTOR = 240  # Фактор для межстрочного интервала
    CHARS_PER_PAGE = 1800  # Примерное количество символов на странице
    
    # Маппинг выравниваний python-docx в строковые значения
    ALIGNMENT_MAP = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
    }
    
    # Стили заголовков по умолчанию (dict для надёжности)
    HEADING_STYLES = {
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Heading 4": 4,
        "Heading 5": 5,
        "Heading 6": 6,
    }
    
    def __init__(self):
        """Инициализация парсера."""
        self._current_page = 1
        self._page_breaks_count = 0
        self._table_counter = 0
        self._figure_counter = 0
        self._logger = None
    
    def parse(self, file_path: str | Path) -> DocumentContent:
        """
        Распарсить DOCX файл и вернуть структурированное содержимое.
        
        Args:
            file_path: Путь к DOCX файлу
            
        Returns:
            DocumentContent: Структурированное содержимое документа
            
        Raises:
            FileNotFoundError: Файл не найден
            ValueError: Ошибка при парсинге файла
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        if path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(f"Неподдерживаемый формат: {path.suffix}")
        
        try:
            doc = Document(str(path))
        except Exception as e:
            raise ValueError(f"Ошибка при открытии DOCX файла: {e}")
        
        content = DocumentContent(
            file_path=str(path),
            file_name=path.name
        )
        
        # Сброс счётчиков
        self._current_page = 1
        self._page_breaks_count = 0
        self._table_counter = 0
        self._figure_counter = 0
        
        # Парсинг элементов документа
        self._parse_sections(doc, content)
        self._parse_paragraphs(doc, content)
        self._parse_tables(doc, content)
        self._parse_figures(doc, content)
        self._parse_formulas(doc, content)
        self._parse_references(content)
        self._parse_bibliography(content)
        self._parse_appendices(content)
        self._parse_lists(content)
        
        # Подсчёт статистики
        self._calculate_statistics(content)
        
        # Подсчёт страниц (приблизительный)
        content.total_pages = self._estimate_pages(content)
        
        return content
    
    def _parse_sections(self, doc: Document, content: DocumentContent) -> None:
        """
        Распарсить секции документа для получения метаданных о полях.
        
        Args:
            doc: Объект документа python-docx
            content: Объект содержимого для сохранения метаданных
        """
        if doc.sections:
            first_section = doc.sections[0]
            content.metadata['margins'] = {
                'left_cm': float(first_section.left_margin.cm) if first_section.left_margin else None,
                'right_cm': float(first_section.right_margin.cm) if first_section.right_margin else None,
                'top_cm': float(first_section.top_margin.cm) if first_section.top_margin else None,
                'bottom_cm': float(first_section.bottom_margin.cm) if first_section.bottom_margin else None,
            }
    
    def _parse_paragraphs(self, doc: Document, content: DocumentContent) -> None:
        """
        Распарсить все абзацы документа.
        
        Args:
            doc: Объект документа python-docx
            content: Объект содержимого для сохранения абзацев и заголовков
        """
        for para in doc.paragraphs:
            # Проверка на разрыв страницы
            if self._has_page_break(para):
                self._current_page += 1
                self._page_breaks_count += 1
            
            # Получение стиля абзаца
            style_name = para.style.name if para.style else None
            
            # Проверка на заголовок
            if style_name and style_name in self.HEADING_STYLES:
                heading = self._parse_heading(para, style_name)
                content.headings.append(heading)
            else:
                paragraph = self._parse_paragraph(para, style_name)
                content.paragraphs.append(paragraph)
    
    def _parse_paragraph(self, para, style_name: Optional[str]) -> Paragraph:
        """
        Распарсить обычный абзац.
        
        Args:
            para: Объект абзаца python-docx
            style_name: Имя стиля абзаца
            
        Returns:
            Paragraph: Структурированный абзац
        """
        text = para.text.strip()
        
        # Получение форматирования из первого запуска (run)
        font_info = self._get_font_info(para)
        
        # Получение выравнивания
        alignment = None
        if para.alignment is not None:
            alignment = self.ALIGNMENT_MAP.get(para.alignment, "left")
        
        # Получение интервалов и отступов
        line_spacing = self._get_line_spacing(para)
        indent_left = self._get_indent_left(para)
        indent_right = self._get_indent_right(para)
        indent_first_line = self._get_indent_first_line(para)
        
        return Paragraph(
            text=text,
            style_name=style_name,
            font_size=font_info.get('size'),
            font_name=font_info.get('name'),
            is_bold=font_info.get('bold', False),
            is_italic=font_info.get('italic', False),
            alignment=alignment,
            line_spacing=line_spacing,
            indent_left=indent_left,
            indent_right=indent_right,
            indent_first_line=indent_first_line,
            page_number=self._current_page
        )
    
    def _parse_heading(self, para, style_name: str) -> Heading:
        """
        Распарсить заголовок.
        
        Args:
            para: Объект абзаца python-docx
            style_name: Имя стиля заголовка
            
        Returns:
            Heading: Структурированный заголовок
        """
        # Определение уровня заголовка через dict lookup (надёжнее enumerate)
        level = self.HEADING_STYLES.get(style_name, 1)
        
        text = para.text.strip()
        
        # Попытка извлечь нумерацию из начала заголовка
        numbering = self._extract_heading_numbering(text)
        
        # Получение форматирования
        font_info = self._get_font_info(para)
        
        # Получение выравнивания
        alignment = None
        if para.alignment is not None:
            alignment = self.ALIGNMENT_MAP.get(para.alignment, "left")
        
        return Heading(
            text=text,
            style_name=style_name,
            level=level,
            numbering=numbering,
            font_size=font_info.get('size'),
            font_name=font_info.get('name'),
            is_bold=font_info.get('bold', False),
            is_italic=font_info.get('italic', False),
            alignment=alignment,
            page_number=self._current_page
        )
    
    def _parse_tables(self, doc: Document, content: DocumentContent) -> None:
        """
        Распарсить все таблицы документа.
        
        Args:
            doc: Объект документа python-docx
            content: Объект содержимого для сохранения таблиц
        """
        for table in doc.tables:
            self._table_counter += 1
            
            rows = len(table.rows)
            columns = len(table.columns)
            
            # Извлечение содержимого ячеек с созданием TableCell объектов
            cells = []
            raw_cells = []
            for row in table.rows:
                row_data = []
                raw_row = []
                for cell in row.cells:
                    raw_text = cell.text.strip()
                    raw_row.append(raw_text)
                    cell_obj = TableCell(
                        text=raw_text,
                        paragraph_count=len(cell.paragraphs)
                    )
                    row_data.append(cell_obj)
                cells.append(row_data)
                raw_cells.append(raw_row)
            
            # Поиск подписи таблицы (обычно абзац перед или после таблицы)
            caption = None
            caption_position = "top"
            
            table_obj = Table(
                rows=rows,
                columns=columns,
                caption=caption,
                caption_position=caption_position,
                page_number=self._current_page,
                table_number=self._table_counter,
                cells=cells,
                raw_cells=raw_cells
            )
            content.tables.append(table_obj)
    
    def _parse_references(self, content: DocumentContent) -> None:
        """
        Распознать библиографические ссылки в тексте.
        
        Поддерживаемые форматы:
        - [1], [1; 2; 3] - простые ссылки
        - [1, с. 5] - ссылка с указанием страницы
        - [там же, с. 5] - ссылка на тот же источник
        
        Args:
            content: Объект содержимого документа
        """
        import re
        
        # Паттерны для ссылок (от более специфичных к менее)
        patterns = [
            (r'\[там же,\s*с\.\s*(\d+)\]', 'same_as'),  # [там же, с. 5]
            (r'\[(\d+),\s*с\.\s*(\d+)\]', 'page'),  # [1, с. 5]
            (r'\[\s*(\d+(?:\s*[;,]\s*\d+)*)\s*\]', 'multiple'),  # [1; 2; 3] или [1, 2]
        ]
        
        for paragraph in content.paragraphs:
            found_refs = set()  # Для предотвращения дублирования
            
            for pattern, pattern_type in patterns:
                matches = re.finditer(pattern, paragraph.text)
                for match in matches:
                    ref_text = match.group(0)
                    
                    # Пропускаем уже найденные ссылки
                    if ref_text in found_refs:
                        continue
                    found_refs.add(ref_text)
                    
                    # Извлечение номеров и страниц
                    numbers = []
                    cited_page = None
                    ref_type = ReferenceType.BRACKETED
                    
                    if pattern_type == 'same_as':
                        ref_type = ReferenceType.SAME_AS
                        cited_page = int(match.group(1))
                    elif pattern_type == 'page':
                        try:
                            numbers = [int(match.group(1))]
                            cited_page = int(match.group(2))
                        except ValueError:
                            pass
                    elif pattern_type == 'multiple':
                        try:
                            first_group = match.group(1)
                            # Разделители могут быть ; или ,
                            numbers = [int(x.strip()) for x in re.split(r'[;,]', first_group)]
                        except ValueError:
                            pass
                    
                    reference = Reference(
                        text=ref_text,
                        reference_type="gost",
                        ref_type=ref_type,
                        position="inline",
                        numbers=numbers,
                        number=numbers[0] if numbers else None,  # Для обратной совместимости
                        cited_page=cited_page,
                        page_number=paragraph.page_number,
                        start_index=match.start(),
                        end_index=match.end()
                    )
                    content.references.append(reference)
    
    def _get_font_info(self, para) -> Dict[str, Any]:
        """
        Получить информацию о шрифте из абзаца.
        
        Args:
            para: Объект абзаца python-docx
            
        Returns:
            Dict: Информация о шрифте (name, size, bold, italic)
        """
        font_info = {
            'name': None,
            'size': None,
            'bold': False,
            'italic': False
        }
        
        # Если есть запуски (runs), берём информацию из первого непустого
        if para.runs:
            for run in para.runs:
                if run.text.strip():
                    if run.font:
                        if run.font.name:
                            font_info['name'] = run.font.name
                        if run.font.size is not None:
                            try:
                                font_info['size'] = float(run.font.size.pt)
                            except (TypeError, AttributeError):
                                pass
                        if run.font.bold is not None:
                            font_info['bold'] = bool(run.font.bold)
                        if run.font.italic is not None:
                            font_info['italic'] = bool(run.font.italic)
                    break
        
        # Если в стилях есть информация о шрифте
        if para.style and hasattr(para.style, 'font') and para.style.font:
            style_font = para.style.font
            if font_info['name'] is None and style_font.name:
                font_info['name'] = style_font.name
            if font_info['size'] is None and style_font.size is not None:
                try:
                    font_info['size'] = float(style_font.size.pt)
                except (TypeError, AttributeError):
                    pass
        
        return font_info
    
    def _get_line_spacing(self, para) -> Optional[float]:
        """
        Получить межстрочный интервал абзаца.
        
        Args:
            para: Объект абзаца python-docx
            
        Returns:
            float: Межстрочный интервал (например, 1.5) или None
        """
        # Попытка получить из XML элемента
        try:
            pPr = para._element.pPr
            if pPr is not None:
                spacing = pPr.find(qn('w:spacing'))
                if spacing is not None:
                    line_rule = spacing.get(qn('w:lineRule'), '')
                    line_val = spacing.get(qn('w:val'))
                    
                    if line_val:
                        # Интерпретация значения
                        if line_rule == 'auto':
                            # Автоматический интервал (обычно 1.0 или 1.15)
                            return 1.0
                        elif line_rule == 'exact':
                            # Точное значение в twips (1/1440 дюйма)
                            # Конвертация в пункты и сравнение с размером шрифта
                            try:
                                twips = int(line_val)
                                # Приблизительная оценка
                                return twips / 240  # грубая эвристика
                            except ValueError:
                                pass
                        elif line_rule == 'atLeast':
                            # Минимальный интервал
                            try:
                                twips = int(line_val)
                                return max(1.0, twips / 240)
                            except ValueError:
                                pass
                        else:
                            # Множитель
                            try:
                                return float(line_val) / 240
                            except ValueError:
                                pass
        except Exception:
            pass
        
        return None
    
    def _get_indent_left(self, para) -> float:
        """
        Получить левый отступ абзаца в пунктах.
        
        Args:
            para: Объект абзаца python-docx
            
        Returns:
            float: Левый отступ в пунктах (0.0 если не указан)
        """
        try:
            pPr = para._element.pPr
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    left = ind.get(qn('w:left'))
                    if left:
                        # Конвертация из twips в пункты (1 пункт = 20 twips)
                        return float(left) / 20
        except Exception:
            pass
        
        return 0.0
    
    def _get_indent_right(self, para) -> float:
        """
        Получить правый отступ абзаца в пунктах.
        
        Args:
            para: Объект абзаца python-docx
            
        Returns:
            float: Правый отступ в пунктах (0.0 если не указан)
        """
        try:
            pPr = para._element.pPr
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    right = ind.get(qn('w:right'))
                    if right:
                        return float(right) / 20
        except Exception:
            pass
        
        return 0.0
    
    def _get_indent_first_line(self, para) -> float:
        """
        Получить отступ первой строки абзаца в пунктах.
        
        Args:
            para: Объект абзаца python-docx
            
        Returns:
            float: Отступ первой строки в пунктах (0.0 если не указан)
        """
        try:
            pPr = para._element.pPr
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    first_line = ind.get(qn('w:firstLine'))
                    if first_line:
                        # Конвертация из twips в пункты
                        return float(first_line) / 20
        except Exception:
            pass
        
        return 0.0
    
    def _has_page_break(self, para) -> bool:
        """
        Проверить наличие разрыва страницы в абзаце.
        
        Args:
            para: Объект абзаца python-docx
            
        Returns:
            bool: True если есть разрыв страницы
        """
        try:
            # Проверка на pageBreakBefore в свойствах абзаца
            if para.paragraph_format and para.paragraph_format.page_break_before:
                return True
            
            # Проверка на явный разрыв страницы в элементе
            for element in para._element:
                tag = element.tag
                if tag.endswith('pageBreak') or (tag.endswith('br') and 'page' in str(element.attrib)):
                    return True
                
                # Проверка на w:br с типом page
                if 'br' in tag:
                    br_type = element.get(qn('w:type'), '')
                    if br_type == 'page':
                        return True
        except Exception:
            pass
        
        return False
    
    def _extract_heading_numbering(self, text: str) -> Optional[str]:
        """
        Извлечь нумерацию из текста заголовка.
        
        Args:
            text: Текст заголовка
            
        Returns:
            str: Нумерация (например, "1.", "1.1") или None
        """
        # Паттерны для различных форматов нумерации (от более специфичных к менее)
        patterns = [
            r'^(Глава\s+\d+)(?:\s|\.|$)',  # Глава 1, Глава 1.
            r'^(\d+\.\d+\.\d+)\.?\s',  # 1.1.1., 1.1.1 
            r'^(\d+\.\d+)\.',  # 1.1.
            r'^(\d+)\.',  # 1.
            r'^([IVX]+)\.',  # Римские цифры
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                return match.group(1)
        
        return None
    
    def _parse_figures(self, doc: Document, content: DocumentContent) -> None:
        """
        Распарсить все рисунки/изображения в документе.
        
        Извлекает inline shapes и пытается найти подписи вида "Рис. N. ..."
        
        Args:
            doc: Объект документа python-docx
            content: Объект содержимого для сохранения рисунков
        """
        # Парсинг inline shapes (изображения вне таблиц)
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    self._figure_counter += 1
                    
                    # Попытка найти подпись рядом с изображением
                    caption = None
                    description = ""
                    
                    # Ищем абзацы с подписями вида "Рис. 1. Описание"
                    for para in doc.paragraphs:
                        match = re.match(r'Рис\.\s*(\d+)\.\s*(.+)', para.text.strip())
                        if match:
                            fig_num = int(match.group(1))
                            description = match.group(2)
                            caption = para.text.strip()
                            
                            figure = Figure(
                                caption=caption,
                                figure_number=fig_num,
                                description=description,
                                page_number=self._current_page,
                                figure_type="image"
                            )
                            content.figures.append(figure)
                            break
                    
                    # Если не нашли стандартную подпись, создаём без неё
                    if not caption:
                        figure = Figure(
                            figure_number=self._figure_counter,
                            page_number=self._current_page,
                            figure_type="image"
                        )
                        content.figures.append(figure)
        except Exception:
            # Если нет изображений или ошибка доступа к rels
            pass
    
    def _parse_formulas(self, doc: Document, content: DocumentContent) -> None:
        """
        Распарсить математические формулы в формате OMML.
        
        Ищет элементы <m:oMath> в XML структуре документа.
        
        Args:
            doc: Объект документа python-docx
            content: Объект содержимого для сохранения формул
        """
        from docx.oxml.ns import qn
        
        # Проход по всем абзацам в поисках формул
        for para in doc.paragraphs:
            # Поиск элементов формул в XML
            formula_elements = para._element.findall('.//' + qn('m:oMath'))
            
            for formula_elem in formula_elements:
                omml_xml = formula_elem.xml if hasattr(formula_elem, 'xml') else str(formula_elem)
                
                # Попытка найти номер формулы справа (обычно в том же абзаце)
                formula_number = None
                text_parts = para.text.split('\t')
                if len(text_parts) > 1:
                    # Номер формулы может быть в конце строки после таб
                    potential_num = text_parts[-1].strip()
                    if re.match(r'\(\d+\)$', potential_num):
                        formula_number = potential_num
                
                formula = Formula(
                    omml_xml=omml_xml,
                    formula_number=formula_number,
                    page_number=self._current_page
                )
                content.formulas.append(formula)
    
    def _parse_bibliography(self, content: DocumentContent) -> None:
        """
        Распарсить список литературы.
        
        Находит раздел "Список литературы" или "Библиографический список"
        и извлекает отдельные записи.
        
        Args:
            content: Объект содержимого документа
        """
        import re
        
        bibliography_started = False
        current_index = 0
        current_entry_lines = []
        
        for i, para in enumerate(content.paragraphs):
            text = para.text.strip()
            
            # Проверка на заголовок раздела библиографии
            if re.search(r'(список\s+литературы|библиографический\s+список|литература)', 
                        text, re.IGNORECASE):
                bibliography_started = True
                continue
            
            if bibliography_started and text:
                # Проверка на начало новой записи (номер в начале)
                match = re.match(r'^(\d+)\.\s+(.+)', text)
                if match:
                    # Сохранение предыдущей записи
                    if current_index > 0 and current_entry_lines:
                        self._create_bibliography_entry(
                            current_index, 
                            current_entry_lines, 
                            content
                        )
                    
                    # Начало новой записи
                    current_index = int(match.group(1))
                    current_entry_lines = [match.group(2)]
                elif current_index > 0:
                    # Продолжение текущей записи
                    current_entry_lines.append(text)
        
        # Сохранение последней записи
        if current_index > 0 and current_entry_lines:
            self._create_bibliography_entry(current_index, current_entry_lines, content)
    
    def _create_bibliography_entry(self, index: int, lines: List[str], 
                                    content: DocumentContent) -> None:
        """
        Создать запись библиографии из строк текста.
        
        Args:
            index: Порядковый номер записи
            lines: Строки текста записи
            content: Объект содержимого для добавления записи
        """
        import re
        
        full_text = ' '.join(lines)
        entry = BibliographyEntry(
            index=index,
            raw_text=full_text,
            title=full_text
        )
        
        # Попытка извлечь авторов (первые слова до точки)
        authors_match = re.match(r'^([А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ]\.?\s*[А-ЯЁ][а-яё]+)*).*?[.,]', full_text)
        if authors_match:
            authors_str = authors_match.group(1)
            entry.authors = [a.strip() for a in re.split(r'[;,]', authors_str)]
        
        # Попытка извлечь URL
        url_match = re.search(r'(https?://[^\s]+)', full_text)
        if url_match:
            entry.url = url_match.group(1)
        
        # Попытка извлечь дату обращения
        access_match = re.search(r'\(обращено:\s*(\d{2}\.\d{2}\.\d{4})\)', full_text, re.IGNORECASE)
        if access_match:
            entry.access_date = access_match.group(1)
        
        content.bibliography.append(entry)
    
    def _parse_appendices(self, content: DocumentContent) -> None:
        """
        Распарсить приложения к документу.
        
        Находит заголовки вида "Приложение А", "Приложение Б" и т.д.
        
        Args:
            content: Объект содержимого документа
        """
        import re
        
        current_appendix = None
        appendix_content = []
        
        for heading in content.headings:
            text = heading.text.strip()
            match = re.match(r'Приложение\s+([А-ЯЁ])\s*[.:]?\s*(.*)', text, re.IGNORECASE)
            
            if match:
                # Сохранение предыдущего приложения
                if current_appendix:
                    current_appendix.content = appendix_content
                    content.appendices.append(current_appendix)
                
                # Начало нового приложения
                letter = match.group(1)
                title = match.group(2).strip()
                
                current_appendix = Appendix(
                    letter=letter,
                    title=title,
                    page_start=heading.page_number
                )
                appendix_content = []
        
        # Сохранение последнего приложения
        if current_appendix:
            current_appendix.content = appendix_content
            content.appendices.append(current_appendix)
    
    def _parse_lists(self, content: DocumentContent) -> None:
        """
        Распознать многоуровневые списки в абзацах.
        
        Определяет маркированные и нумерованные списки.
        
        Args:
            content: Объект содержимого для обновления информации о списках
        """
        import re
        
        for para in content.paragraphs:
            text = para.text.strip()
            
            # Проверка на маркированный список
            if re.match(r'^[•●○■□▪▫–—-]\s+', text):
                para.is_list_item = True
                para.list_type = 'bullet'
                # Определение уровня по отступу
                if para.indent_left > 0:
                    para.list_level = min(5, int(para.indent_left / 36) + 1)
                else:
                    para.list_level = 1
            
            # Проверка на нумерованный список
            elif re.match(r'^(\d+|[IVXivx]+|[а-яёA-Z])[\.\)]\s+', text):
                para.is_list_item = True
                para.list_type = 'number'
                if para.indent_left > 0:
                    para.list_level = min(5, int(para.indent_left / 36) + 1)
                else:
                    para.list_level = 1
    
    def _calculate_statistics(self, content: DocumentContent) -> None:
        """
        Подсчитать статистику по документу.
        
        Args:
            content: Объект содержимого для обновления статистики
        """
        # Подсчёт символов
        all_text = ' '.join([p.text for p in content.paragraphs])
        all_text += ' '.join([h.text for h in content.headings])
        
        content.char_count_with_spaces = len(all_text)
        content.char_count_no_spaces = len(all_text.replace(' ', ''))
        content.word_count = len(all_text.split())
    
    def _estimate_pages(self, content: DocumentContent) -> int:
        """
        Оценить количество страниц в документе.
        
        Args:
            content: Объект содержимого документа
            
        Returns:
            int: Примерное количество страниц
        """
        # Более точная эвристика с учётом средней длины строки
        total_chars = content.char_count_no_spaces
        
        # Базовая оценка по символам
        estimated_pages = max(1, total_chars // self.CHARS_PER_PAGE)
        
        # Учёт таблиц и рисунков
        estimated_pages += len(content.tables) // 3
        estimated_pages += len(content.figures) // 3
        estimated_pages += len(content.formulas) // 5
        
        # Берём максимум из подсчитанных разрывов страниц и оценки
        return max(self._current_page, estimated_pages)


def parse_docx(file_path: str | Path) -> DocumentContent:
    """
    Удобная функция для парсинга DOCX файла.
    
    Args:
        file_path: Путь к DOCX файлу
        
    Returns:
        DocumentContent: Структурированное содержимое документа
    """
    parser = DOCXParser()
    return parser.parse(file_path)
